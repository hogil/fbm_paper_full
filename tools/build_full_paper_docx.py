from __future__ import annotations

import argparse
from datetime import datetime
import re
import shutil
import zipfile
from xml.etree import ElementTree as ET
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[1]
INPUT = ROOT / "full_paper_rev001_codex.md"
OUTPUT = ROOT / "full_paper_rev001_codex.docx"

FONT_KO = "바탕체"
FONT_EN = "Times New Roman"
CORE_TIMESTAMP = datetime(2026, 6, 6, 0, 0, 0)


def set_run_font(run, size_pt: float, bold: bool = False, italic: bool = False):
    run.font.name = FONT_EN
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.italic = italic
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_KO)
    run._element.rPr.rFonts.set(qn("w:ascii"), FONT_EN)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_EN)


def set_paragraph_format(p, *, align=WD_ALIGN_PARAGRAPH.JUSTIFY, before=0, after=3):
    p.alignment = align
    pf = p.paragraph_format
    pf.line_spacing = 1.0
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)


def set_page(section):
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)
    section.header_distance = Cm(1.0)
    section.footer_distance = Cm(0.5)


def set_columns(section, count: int = 2, space_cm: float = 0.5):
    sect_pr = section._sectPr
    cols = sect_pr.xpath("./w:cols")
    if cols:
        cols = cols[0]
    else:
        cols = OxmlElement("w:cols")
        sect_pr.append(cols)
    cols.set(qn("w:num"), str(count))
    cols.set(qn("w:space"), str(int(space_cm * 567)))


def set_one_column(section):
    sect_pr = section._sectPr
    cols = sect_pr.xpath("./w:cols")
    if cols:
        cols[0].set(qn("w:num"), "1")


def add_page_number(paragraph):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr)
    run._r.append(fld_char2)


def configure_header_footer(section, first_page_affiliation: str | None = None):
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    header_p = section.header.paragraphs[0]
    header_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header_p.text = ""
    r = header_p.add_run("Samsung Best Paper Awards")
    set_run_font(r, 12, bold=True)

    footer_p = section.footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_p.text = ""
    add_page_number(footer_p)

    if first_page_affiliation:
        section.different_first_page_header_footer = True
        section.first_page_header.is_linked_to_previous = False
        section.first_page_footer.is_linked_to_previous = False
        first_header = section.first_page_header.paragraphs[0]
        first_header.alignment = WD_ALIGN_PARAGRAPH.LEFT
        first_header.text = ""
        r = first_header.add_run("Samsung Best Paper Awards")
        set_run_font(r, 12, bold=True)

        first_footer = section.first_page_footer
        fp0 = first_footer.paragraphs[0]
        fp0.alignment = WD_ALIGN_PARAGRAPH.LEFT
        fp0.text = ""
        r = fp0.add_run(first_page_affiliation)
        set_run_font(r, 9)
        fp1 = first_footer.add_paragraph()
        fp1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_page_number(fp1)


def configure_visible_first_page_footer(section, affiliation: str):
    section.footer.is_linked_to_previous = False
    footer = section.footer
    fp0 = footer.paragraphs[0]
    fp0.alignment = WD_ALIGN_PARAGRAPH.LEFT
    fp0.text = ""
    r = fp0.add_run(affiliation)
    set_run_font(r, 9)
    fp1 = footer.add_paragraph()
    fp1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_number(fp1)


def clear_cell_borders(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:" + edge
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        element.set(qn("w:val"), "nil")


def set_cell_border(cell, edge: str, val="single", sz="8", color="000000"):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    element = tc_borders.find(qn("w:" + edge))
    if element is None:
        element = OxmlElement("w:" + edge)
        tc_borders.append(element)
    element.set(qn("w:val"), val)
    element.set(qn("w:sz"), sz)
    element.set(qn("w:space"), "0")
    element.set(qn("w:color"), color)


def add_text_paragraph(doc, text: str, size=10, bold=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY, after=3):
    p = doc.add_paragraph()
    set_paragraph_format(p, align=align, after=after)
    r = p.add_run(text)
    set_run_font(r, size, bold=bold)
    return p


def set_row_cant_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def add_inline_markdown_paragraph(doc, text: str, size=10, align=WD_ALIGN_PARAGRAPH.LEFT, after=3):
    p = doc.add_paragraph()
    set_paragraph_format(p, align=align, after=after)
    pattern = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*)")
    pos = 0
    for match in pattern.finditer(text):
        if match.start() > pos:
            r = p.add_run(text[pos:match.start()])
            set_run_font(r, size)
        token = match.group(0)
        if token.startswith("**"):
            r = p.add_run(token[2:-2])
            set_run_font(r, size, bold=True)
        else:
            r = p.add_run(token[1:-1])
            set_run_font(r, size, italic=True)
        pos = match.end()
    if pos < len(text):
        r = p.add_run(text[pos:])
        set_run_font(r, size)
    return p


def add_heading(doc, text: str):
    p = doc.add_paragraph()
    set_paragraph_format(p, align=WD_ALIGN_PARAGRAPH.LEFT, before=8, after=4)
    r = p.add_run(text)
    set_run_font(r, 12, bold=True)


def add_subheading(doc, text: str):
    p = doc.add_paragraph()
    set_paragraph_format(p, align=WD_ALIGN_PARAGRAPH.LEFT, before=5, after=2)
    r = p.add_run(text)
    set_run_font(r, 10, bold=True)


def strip_bold_marks(text: str) -> str:
    if text.startswith("**") and text.endswith("**"):
        return text[2:-2]
    return text.replace("**", "")


def split_table_row(line: str) -> list[str]:
    line = line.strip()
    if line.startswith("|"):
        line = line[1:]
    if line.endswith("|"):
        line = line[:-1]
    return [c.strip() for c in line.split("|")]


def add_markdown_table(doc, caption: str, rows: list[str]):
    caption_p = add_text_paragraph(doc, caption, size=9, align=WD_ALIGN_PARAGRAPH.LEFT, after=2)
    caption_p.paragraph_format.keep_with_next = True
    headers = split_table_row(rows[0])
    body = [split_table_row(r) for r in rows[2:]]
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    set_row_cant_split(table.rows[0])

    for idx, text in enumerate(headers):
        cell = table.rows[0].cells[idx]
        clear_cell_borders(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        set_paragraph_format(p, align=WD_ALIGN_PARAGRAPH.CENTER, after=0)
        r = p.add_run(text)
        set_run_font(r, 9, bold=True)
        set_cell_border(cell, "top", sz="12")
        set_cell_border(cell, "bottom", sz="8")

    for row_values in body:
        row = table.add_row()
        set_row_cant_split(row)
        for idx, text in enumerate(row_values):
            cell = row.cells[idx]
            clear_cell_borders(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            set_paragraph_format(p, align=WD_ALIGN_PARAGRAPH.CENTER, after=0)
            r = p.add_run(text)
            set_run_font(r, 9)

    for cell in table.rows[-1].cells:
        set_cell_border(cell, "bottom", sz="12")

    doc.add_paragraph()


def add_figure_label_row(doc, labels: list[str]):
    if not labels:
        return
    p = doc.add_paragraph()
    set_paragraph_format(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY, after=0)
    p.paragraph_format.keep_with_next = True
    stops = p.paragraph_format.tab_stops
    count = len(labels)
    for idx in range(count):
        pos = 8.0 * (idx + 0.5) / count
        stops.add_tab_stop(Cm(pos), WD_TAB_ALIGNMENT.CENTER)
    r = p.add_run("\t" + "\t".join(labels))
    set_run_font(r, 9)


def add_figure(doc, image_path: str, caption: str, labels: list[str] | None = None):
    add_figure_label_row(doc, labels or [])
    p = doc.add_paragraph()
    set_paragraph_format(p, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
    img = ROOT / image_path
    if img.exists():
        p.add_run().add_picture(str(img), width=Cm(8.0))
    else:
        r = p.add_run(f"[Missing figure: {image_path}]")
        set_run_font(r, 9, bold=True)
    cap = doc.add_paragraph()
    set_paragraph_format(cap, align=WD_ALIGN_PARAGRAPH.CENTER, after=4)
    r = cap.add_run(caption)
    set_run_font(r, 9)


def set_core_properties(doc, title: str, authors: str):
    props = doc.core_properties
    props.title = title
    props.subject = "Samsung Best Paper Awards full paper"
    props.author = authors.replace("¹", "")
    props.last_modified_by = authors.replace("¹", "")
    props.keywords = "Failbit Map; Known Defect; Unknown Defect; Object-id Map; HDBSCAN"
    props.comments = "Samsung Best Paper Awards full-paper submission candidate."
    props.category = "Samsung Best Paper Awards"
    props.revision = 1
    props.created = CORE_TIMESTAMP
    props.modified = CORE_TIMESTAMP


def sanitize_app_properties(docx_path: Path):
    app_name = "docProps/app.xml"
    ns = {
        "ep": "http://schemas.openxmlformats.org/officeDocument/2006/extended-properties",
        "vt": "http://schemas.openxmlformats.org/officeDocument/2006/docPropsVTypes",
    }
    tmp_path = docx_path.with_suffix(docx_path.suffix + ".tmp")
    with zipfile.ZipFile(docx_path, "r") as src, zipfile.ZipFile(tmp_path, "w", zipfile.ZIP_DEFLATED) as dst:
        for item in src.infolist():
            data = src.read(item.filename)
            if item.filename == app_name:
                root = ET.fromstring(data)
                app = root.find("ep:Application", ns)
                if app is not None:
                    app.text = "Microsoft Word"
                company = root.find("ep:Company", ns)
                if company is not None:
                    company.text = "Samsung Electronics"
                data = ET.tostring(root, encoding="utf-8", xml_declaration=True)
            dst.writestr(item, data)
    shutil.move(str(tmp_path), str(docx_path))


def build(input_path: Path = INPUT, output_path: Path = OUTPUT):
    lines = input_path.read_text(encoding="utf-8").splitlines()
    doc = Document()
    set_page(doc.sections[0])
    set_one_column(doc.sections[0])
    configure_header_footer(doc.sections[0])

    normal = doc.styles["Normal"]
    normal.font.name = FONT_EN
    normal.font.size = Pt(10)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_KO)
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT_EN)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_EN)

    idx = 0
    title = lines[idx].lstrip("# ").strip()
    p = add_text_paragraph(doc, title, size=20, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT, after=6)
    idx += 2

    authors = lines[idx].strip()
    set_core_properties(doc, title, authors)
    add_text_paragraph(doc, authors, size=12, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT, after=2)
    idx += 2

    affiliation = lines[idx].strip()
    idx += 2
    configure_visible_first_page_footer(doc.sections[0], affiliation)

    if lines[idx].strip() == "**Abstract**":
        add_text_paragraph(doc, "Abstract", size=10, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT, after=2)
        idx += 2
        abstract = strip_bold_marks(lines[idx].strip())
        add_text_paragraph(doc, abstract, size=10, bold=True, align=WD_ALIGN_PARAGRAPH.JUSTIFY, after=5)
        idx += 2

    if idx < len(lines) and lines[idx].startswith("Keywords:"):
        add_text_paragraph(doc, lines[idx], size=9, align=WD_ALIGN_PARAGRAPH.LEFT, after=8)
        idx += 1

    body_section = doc.add_section(WD_SECTION.CONTINUOUS)
    set_page(body_section)
    set_columns(body_section, 2, 0.5)
    configure_header_footer(body_section, first_page_affiliation=affiliation)

    pending_figure_labels = None
    while idx < len(lines):
        line = lines[idx].strip()
        if not line:
            idx += 1
            continue

        if line == "<!-- column-break -->":
            doc.add_paragraph().add_run().add_break(WD_BREAK.COLUMN)
            idx += 1
            continue

        if line.startswith("## References"):
            add_heading(doc, "References")
            idx += 1
            continue

        if line.startswith("## "):
            add_heading(doc, line[3:].strip())
            idx += 1
            continue

        if line.startswith("### "):
            add_subheading(doc, line[4:].strip())
            idx += 1
            continue

        if line.startswith("Figure labels:"):
            pending_figure_labels = [part.strip() for part in line.split(":", 1)[1].split("|") if part.strip()]
            idx += 1
            continue

        if line.startswith("!["):
            m = re.match(r"!\[[^\]]*\]\(([^)]+)\)", line)
            img_path = m.group(1) if m else ""
            caption = ""
            lookahead = idx + 1
            while lookahead < len(lines) and not lines[lookahead].strip():
                lookahead += 1
            if lookahead < len(lines) and lines[lookahead].strip().startswith("Figure "):
                caption = lines[lookahead].strip()
                idx = lookahead + 1
            else:
                idx += 1
            add_figure(doc, img_path, caption, pending_figure_labels)
            pending_figure_labels = None
            continue

        if line.startswith("Figure "):
            idx += 1
            continue

        if re.match(r"^Table \d+\.", line):
            caption = line
            table_lines = []
            idx += 1
            while idx < len(lines) and not lines[idx].strip():
                idx += 1
            while idx < len(lines) and lines[idx].strip().startswith("|"):
                table_lines.append(lines[idx].strip())
                idx += 1
            if table_lines:
                add_markdown_table(doc, caption, table_lines)
            else:
                add_text_paragraph(doc, caption, size=9, align=WD_ALIGN_PARAGRAPH.LEFT)
            continue

        if line.startswith("[") and re.match(r"^\[\d+\]", line):
            add_inline_markdown_paragraph(doc, line, size=10, align=WD_ALIGN_PARAGRAPH.LEFT, after=3)
            idx += 1
            continue

        add_text_paragraph(doc, strip_bold_marks(line), size=10, align=WD_ALIGN_PARAGRAPH.JUSTIFY, after=3)
        idx += 1

    doc.save(output_path)
    sanitize_app_properties(output_path)
    print(output_path)


def parse_args():
    parser = argparse.ArgumentParser(description="Build Samsung full paper DOCX from markdown.")
    parser.add_argument("input", nargs="?", default=str(INPUT), help="Input markdown path")
    parser.add_argument("output", nargs="?", default=None, help="Output docx path")
    return parser.parse_args()


if __name__ == "__main__":
    args = parse_args()
    input_path = Path(args.input)
    if not input_path.is_absolute():
        input_path = ROOT / input_path
    if args.output is None:
        output_path = input_path.with_suffix(".docx")
    else:
        output_path = Path(args.output)
        if not output_path.is_absolute():
            output_path = ROOT / output_path
    build(input_path, output_path)
