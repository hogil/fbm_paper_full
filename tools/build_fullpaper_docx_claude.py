"""Claude full paper (.md) -> Samsung Best Paper Awards .docx.

codex 빌더(build_full_paper_docx.py)의 양식 함수(A4·여백·머리글/바닥글·2단·바탕체/Times·
글자크기·3선표·캡션)를 import 재사용하고(수정하지 않음), claude md 포맷
(영문 부제 / <sup> / **Table**·**Figure** / $$math / ```code / > blockquote / [^n] 각주)을
파싱하는 부분만 새로 둔다.

usage:
  python tools/build_fullpaper_docx_claude.py [input.md] [output.docx]
  인자 없으면 full_paper_rev{최대}_claude.md 를 빌드한다.
"""
from __future__ import annotations

import os
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

THIS = Path(__file__).resolve()
ROOT = THIS.parents[1]
sys.path.insert(0, str(THIS.parent))
import build_full_paper_docx as cx  # noqa: E402  (read-only reuse of codex format helpers)

FONT_KO, FONT_EN, MONO = cx.FONT_KO, cx.FONT_EN, "Consolas"
J, L, C = WD_ALIGN_PARAGRAPH.JUSTIFY, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER
FORCE_TWO_COLUMN = os.environ.get("FULLPAPER_FORCE_TWO_COLUMN", "1") != "0"

INLINE = re.compile(r"(\*\*.+?\*\*|\*[^*]+\*|<sup>.*?</sup>|\[\^\d+\])")
SUP = re.compile(r"<sup>(.*?)</sup>")


def _font(run, size, bold=False, italic=False, name=FONT_EN, sup=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    rpr = run._element.get_or_add_rPr()
    rf = rpr.find(qn("w:rFonts"))
    if rf is None:
        rf = OxmlElement("w:rFonts")
        rpr.insert(0, rf)
    rf.set(qn("w:eastAsia"), FONT_KO)
    rf.set(qn("w:ascii"), name)
    rf.set(qn("w:hAnsi"), name)
    if sup:
        run.font.superscript = True


def _run(p, text, size=10, bold=False, italic=False, name=FONT_EN, sup=False):
    r = p.add_run(text)
    _font(r, size, bold, italic, name, sup)
    return r


def _para(doc, align=J, before=0, after=3, indent=None):
    p = doc.add_paragraph()
    cx.set_paragraph_format(p, align=align, before=before, after=after)
    if indent is not None:
        p.paragraph_format.left_indent = Cm(indent)
    return p


def _inline(p, text, size=10, bold=False):
    pos = 0
    for m in INLINE.finditer(text):
        if m.start() > pos:
            _run(p, text[pos:m.start()], size, bold=bold)
        t = m.group(0)
        if t.startswith("**"):
            _run(p, t[2:-2], size, bold=True)
        elif t.startswith("<sup>"):
            _run(p, SUP.match(t).group(1), size, bold=bold, sup=True)
        elif t.startswith("[^"):
            _run(p, t[2:-1], size, sup=True)
        else:
            _run(p, t[1:-1], size, italic=True, bold=bold)
        pos = m.end()
    if pos < len(text):
        _run(p, text[pos:], size, bold=bold)


def _strip_md(t):
    return t.replace("**", "")


def _figure(doc, path, caption, wide=False):
    p = _para(doc, C, after=2)
    if FORCE_TWO_COLUMN:
        wide = False
    img = ROOT / path
    if not img.exists():
        alt = ROOT / "recommendation" / "figures" / Path(path).name
        if alt.exists():
            img = alt
    if img.exists():
        width_cm = 13.4 if wide else 8.2
        if (not FORCE_TWO_COLUMN and
                Path(path).name in {"fig02_pipeline_architecture.png", "fig06_fcm_pm.png", "fig07_probability_control.png",
                                    "p1_fig_pipeline_arch.png", "p2_fig_fcm_pm.png", "p2_fig_prob_control.png"}):
            width_cm = 17.4
        p.add_run().add_picture(str(img), width=Cm(width_cm))
    else:
        _run(p, f"[Missing figure: {path}]", 9, bold=True)
    cap = _para(doc, C, after=4)
    _inline(cap, _strip_md(caption), 9)


def _figure_pair(doc, paths, caption, wide=False):
    """Place 2 images side by side (horizontal) in one paragraph, then caption."""
    p = _para(doc, C, after=2)
    each_cm = 4.0
    for k, path in enumerate(paths[:2]):
        img = ROOT / path
        if not img.exists():
            alt = ROOT / "recommendation" / "figures" / Path(path).name
            if alt.exists():
                img = alt
        if img.exists():
            if k:
                p.add_run("   ")
            p.add_run().add_picture(str(img), width=Cm(each_cm))
        else:
            _run(p, f"[Missing figure: {path}]", 9, bold=True)
    cap = _para(doc, C, after=4)
    _inline(cap, _strip_md(caption), 9)


def _fix_table_width(table, wide=False):
    """Pin table to a fixed layout at column (or full-page) width so autofit cannot
    overflow into the adjacent column/margin (FP Rule A/B gate). Long cells wrap instead."""
    if FORCE_TWO_COLUMN:
        wide = False
    table.autofit = False
    table.allow_autofit = False
    total_cm = 17.8 if wide else 8.6
    ncol = max(1, len(table.columns))
    cw = Cm(total_cm / ncol)
    tblPr = table._tbl.tblPr
    layout = tblPr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tblPr.append(layout)
    layout.set(qn("w:type"), "fixed")
    table.width = Cm(total_cm)
    for col in table.columns:
        col.width = cw
    for row in table.rows:
        for cell in row.cells:
            cell.width = cw


def _bold_matching_rows(table, tokens=("FCM-M", "FCM+margin", "FCM-PM + val_margin")):
    for row in table.rows[1:]:
        row_text = " ".join(cell.text for cell in row.cells)
        if any(tok in row_text for tok in tokens):
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True


def _latex(t):
    t = t.strip().strip("$").strip()
    repl = {
        "\\arg\\max": " argmax ", "\\text{": "", "\\;": " ", "\\,": " ",
        "\\qquad": "     ", "\\quad": "   ", "\\cdot": "·", "}": "", "{": "",
    }
    for a, b in repl.items():
        t = t.replace(a, b)
    return re.sub(r"\s+", " ", t).strip()


def _strip_comments(raw):
    out, in_c = [], False
    for ln in raw:
        s = ln.strip()
        if s in ("<!-- wide-start -->", "<!-- wide-end -->"):
            if not FORCE_TWO_COLUMN:
                out.append(ln)
            continue
        if not in_c and s.startswith("<!--"):
            in_c = not s.endswith("-->")
            continue
        if in_c:
            if s.endswith("-->"):
                in_c = False
            continue
        out.append(ln)
    return out


def build(input_path: Path, output_path: Path):
    lines = _strip_comments(input_path.read_text(encoding="utf-8").splitlines())

    doc = Document()
    s0 = doc.sections[0]
    cx.set_page(s0)
    cx.set_one_column(s0)
    cx.configure_header_footer(s0)
    nm = doc.styles["Normal"]
    nm.font.name = FONT_EN
    nm.font.size = Pt(10)
    nm._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_KO)
    nm._element.rPr.rFonts.set(qn("w:ascii"), FONT_EN)
    nm._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_EN)

    # ── header zone (single column) ──
    title = lines[0].lstrip("#").strip()
    p = _para(doc, L, after=3)
    _run(p, title, 20, bold=True)

    subtitle = author = affiliation = abstract = keywords = None
    affiliations = []
    i = 1
    while i < len(lines):
        s = lines[i].strip()
        if not s or s == "---":
            i += 1
            continue
        if s.startswith("## "):
            break
        if s.startswith("**Abstract**"):
            abstract = s
            i += 1
            continue
        if s.lower().startswith("**keywords") or s.lower().startswith("keywords"):
            keywords = s
            i += 1
            continue
        if subtitle is None and s.startswith("**") and s.endswith("**") and "Abstract" not in s:
            subtitle = _strip_md(s)
            i += 1
            continue
        if author is None and "<sup>" in s and not s.startswith("<sup>"):
            author = s
            i += 1
            continue
        if s.startswith("<sup>"):
            affiliations.append(s)
            if affiliation is None:
                affiliation = s
            i += 1
            continue
        i += 1
    body_start = i

    if subtitle:
        p = _para(doc, L, after=4)
        _run(p, subtitle, 12, bold=True)
    if author:
        p = _para(doc, L, after=2)
        _inline(p, author, 12, bold=True)
    if affiliations:
        aff_plain = "\n".join(re.sub(r"</?sup>", "", a).strip() for a in affiliations)
    else:
        aff_plain = re.sub(r"</?sup>", "", affiliation).strip() if affiliation else ""
    cx.set_core_properties(doc, title, re.sub(r"<sup>.*?</sup>", "", author or "").strip())
    if aff_plain:
        cx.configure_visible_first_page_footer(s0, aff_plain)
    if abstract:
        m = re.match(r"\*\*Abstract\*\*\s*[—\-–]?\s*(.*)", abstract)
        body = _strip_md(m.group(1)) if m else _strip_md(abstract)
        p = _para(doc, J, after=5)
        _run(p, "Abstract - ", 10, bold=True)
        _run(p, body, 10, bold=True)
    if keywords:
        p = _para(doc, L, after=8)
        _inline(p, keywords, 9)

    # ── body (two columns) ──
    bs = doc.add_section(WD_SECTION.CONTINUOUS)
    cx.set_page(bs)
    cx.set_columns(bs, 2, 0.5)
    cx.configure_header_footer(bs, first_page_affiliation=aff_plain or None)

    idx = body_start
    foot_hdr = False
    in_wide = False
    while idx < len(lines):
        s = lines[idx].strip()
        if not s or s == "---":
            idx += 1
            continue

        if s == "<!-- wide-start -->":
            if FORCE_TWO_COLUMN:
                idx += 1
                continue
            ws = doc.add_section(WD_SECTION.CONTINUOUS)
            cx.set_page(ws)
            cx.set_one_column(ws)
            cx.configure_header_footer(ws, first_page_affiliation=aff_plain or None)
            in_wide = True
            idx += 1
            continue

        if s == "<!-- wide-end -->":
            if FORCE_TWO_COLUMN:
                idx += 1
                continue
            ts = doc.add_section(WD_SECTION.CONTINUOUS)
            cx.set_page(ts)
            cx.set_columns(ts, 2, 0.5)
            cx.configure_header_footer(ts, first_page_affiliation=aff_plain or None)
            in_wide = False
            idx += 1
            continue

        if s.startswith("## "):
            cx.add_heading(doc, s[3:].strip())
            idx += 1
            continue
        if s.startswith("#### "):
            cx.add_subheading(doc, s[5:].strip())
            idx += 1
            continue
        if s.startswith("### "):
            cx.add_subheading(doc, s[4:].strip())
            idx += 1
            continue

        if s.startswith("!["):
            paths = re.findall(r"!\[[^\]]*\]\(([^)]+)\)", s)
            cap, la = "", idx + 1
            while la < len(lines) and not lines[la].strip():
                la += 1
            if la < len(lines) and re.match(r"\*?\*?Figure \d+\.", lines[la].strip()):
                cap = lines[la].strip()
                idx = la + 1
            else:
                idx += 1
            if len(paths) >= 2:
                _figure_pair(doc, paths, cap, wide=in_wide)
            else:
                _figure(doc, paths[0] if paths else "", cap, wide=in_wide)
            continue

        if re.match(r"^\*?\*?Figure \d+\.", s):
            idx += 1
            continue

        tbl = _strip_md(s)
        if re.match(r"^Table \d+\.", tbl):
            cap = tbl
            idx += 1
            while idx < len(lines) and not lines[idx].strip():
                idx += 1
            tl = []
            while idx < len(lines) and lines[idx].strip().startswith("|"):
                tl.append(lines[idx].strip().replace("**", ""))
                idx += 1
            if tl:
                cx.add_markdown_table(doc, cap, tl)
                _fix_table_width(doc.tables[-1], wide=in_wide)
                _bold_matching_rows(doc.tables[-1])
            else:
                p = _para(doc, L, after=2)
                _run(p, cap, 9)
            continue

        if s.startswith("$$"):
            block = s
            if not (s.endswith("$$") and len(s) >= 4):
                idx += 1
                while idx < len(lines) and "$$" not in lines[idx]:
                    block += " " + lines[idx].strip()
                    idx += 1
                if idx < len(lines):
                    block += " " + lines[idx].strip()
            p = _para(doc, C, before=2, after=4)
            _run(p, _latex(block), 10, italic=True)
            idx += 1
            continue

        if s.startswith("```"):
            idx += 1
            code = []
            while idx < len(lines) and not lines[idx].strip().startswith("```"):
                code.append(lines[idx])
                idx += 1
            idx += 1
            for cl in code:
                p = _para(doc, L, before=0, after=0)
                _run(p, cl if cl else " ", 8, name=MONO)
            doc.add_paragraph()
            continue

        if s.startswith(">"):
            while idx < len(lines) and lines[idx].strip().startswith(">"):
                b = lines[idx].strip().lstrip(">").strip()
                idx += 1
                if not b:
                    continue
                p = _para(doc, J, after=2, indent=0.4)
                _inline(p, b, 9.5)
            continue

        fm = re.match(r"^\[\^(\d+)\]:\s*(.*)", s)
        if fm:
            if not foot_hdr:
                cx.add_subheading(doc, "Footnotes" if "_en" in input_path.stem.lower() else "각주")
                foot_hdr = True
            p = _para(doc, L, after=2)
            _run(p, f"[{fm.group(1)}] ", 9, bold=True)
            _inline(p, fm.group(2), 9)
            idx += 1
            continue

        if re.match(r"^\[\d+\]\s", s):
            p = _para(doc, L, after=2)
            _inline(p, s, 10)
            idx += 1
            continue

        p = _para(doc, J, after=3)
        _inline(p, s, 10)
        idx += 1

    doc.save(output_path)
    cx.sanitize_app_properties(output_path)
    print(output_path)


def _latest_claude_md():
    cands = []
    for f in os.listdir(ROOT):
        m = re.match(r"full_paper_rev(\d+)_claude\.md$", f)
        if m:
            cands.append((int(m.group(1)), f))
    if not cands:
        raise SystemExit("no full_paper_rev*_claude.md found")
    return ROOT / max(cands)[1]


if __name__ == "__main__":
    args = sys.argv[1:]
    inp = Path(args[0]) if args else _latest_claude_md()
    if not inp.is_absolute():
        inp = ROOT / inp
    out = Path(args[1]) if len(args) > 1 else inp.with_suffix(".docx")
    if not out.is_absolute():
        out = ROOT / out
    build(inp, out)
