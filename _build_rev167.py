#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""rev159: rev158 양식 유지 + 본문 수정본으로 KO/EN DOCX 재생성."""

from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from generate_papers import (
    set_margins, set_two_columns, add_title_block, add_author_block,
    add_abstract_block, add_heading, add_subheading, add_body,
    add_1col_header_break, add_labeled_example_table, add_figure,
    add_table, add_refs, make_yolo_roi_figure,
    TITLE_KO_CODEX, TITLE_EN_CODEX, AUTHORS, AFFIL,
    TABLE_PERF_CLAUDE, REFS_CODEX,
    F_KO, F_EN, PT9, PT10,
)

OUT_DIR = Path(__file__).parent

AUTHORS_EN = "Gildong Hong\u00b9, Cheolsu Kim\u00b9"
AFFIL_EN = "\u00b9 Semiconductor Research Center, Samsung Electronics, Hwaseong-si, Republic of Korea"

ABSTRACT_KO = (
    "Failbit Map은 반도체 EDS Test에서 생성되는 웨이퍼당 약 1,000만 pixel 수준의 초고해상도 데이터로, 불량 패턴 분석의 핵심 자료이다. "
    "그러나 실제 현업에서는 대량의 Failbit Map 조회가 불가능하고, 조회 가능한 일부 Map의 분석조차 엔지니어의 수작업에 의존하고 있다. "
    "본 논문은 이를 해결하기 위해 대량 Failbit Map 데이터 파이프라인을 구축하고, Known 불량은 2-stage supervised classification으로, Unknown 불량은 self-supervised로 검출하는 아키텍처를 구현하였다. "
    "변환 속도는 Cython 적용으로 약 100배 향상되었고, 이미지 용량은 32-color palette-indexed PNG 적용으로 약 75% 절감되었다. "
    "Known 불량 분류는 ConvNeXtV2 기반 1차 분류와 저신뢰 샘플에 대한 ROI(Region of Interest) 기반 YOLO 2차 분류를 결합한 구조로 설계하였으며, weighted F1 0.95를 달성하였다. "
    "Unknown 불량 검출에서는 SimCLR 계열 모델에 대조학습을 적용하였으며, wafer의 zone 기반 불량 해석 특성을 반영하기 위해 grid-structured local sampling을 함께 적용하였다. "
    "양산 5일치 Failbit Map 10,000장을 학습에 사용하고 별도 1일치 2,000장에 적용한 결과 13개 불량 그룹이 검출되었으며, "
    "이 중 7개가 현업 엔지니어 검증에서 실제 불량 그룹으로 판정되어 현업 적용 가능성을 입증하였다."
)

ABSTRACT_EN = (
    "Failbit Maps are ultra-high-resolution data generated in semiconductor EDS tests, with roughly 10 million pixels per wafer, and are a key source for defect-pattern analysis. "
    "However, in practical operations, large-scale retrieval and inspection of Failbit Maps have not been possible, and even the limited analysis that was feasible relied heavily on manual inspection by engineers. "
    "This work addresses that limitation by building a large-scale Failbit Map data pipeline and implementing an architecture in which Known defects are handled by two-stage supervised classification and Unknown defects are discovered by self-supervised learning. "
    "Data transformation speed was improved by about 100x through Cython, while the image storage footprint was reduced by about 75% through a 32-color palette-indexed PNG. "
    "For Known defects, a ConvNeXtV2-based first-stage classifier was combined with a second-stage ROI (Region of Interest)-based YOLO classifier for low-confidence samples, achieving a weighted F1 of 0.95. "
    "For Unknown defects, a SimCLR-family model was trained with contrastive learning using grid-structured local sampling to reflect the zone-based nature of wafer defect interpretation. "
    "After training on 10,000 Failbit Maps collected over five production days, the model was applied to 2,000 maps from an additional day. "
    "It detected 13 defect groups, 7 of which were confirmed by domain engineers as real defect groups, demonstrating practical applicability in production."
)

KEYWORDS_LINE = (
    "Keywords: Failbit Map, Wafer Failure Analysis, ConvNeXtV2, "
    "YOLO, Contrastive Learning, HDBSCAN"
)

# 본문 문단 (KO)
KO_1_1 = (
    "Failbit Map은 EDS Test에서 Memory Cell Block 단위의 불량 정도를 Grade 0부터 7까지로 표현한 데이터이다. "
    "Wafer 1장에는 약 1,000만 개의 block이 존재하므로, 이는 불량의 위치와 형태를 반영하는 초고해상도 데이터이다. "
    "현업의 Measure 기반 분석만으로는 이러한 정보를 확인할 수 없어 Failbit Map의 전수 분석이 필요하다."
)

KO_1_2 = (
    "실제 현업 적용에는 두 가지 제약이 있다. 첫째, 기존 시스템은 설비 Log를 대량 Failbit Map으로 변환하고 저장 및 조회하는 처리 성능이 부족하였다. "
    "설비 Log는 Wafer당 10~50MB 수준이며 전 제품 기준 하루 약 2만 장의 Wafer가 생산되지만, 기존 환경에서는 속도와 메모리 제약으로 대량 처리가 어려웠고 엔지니어가 한 번에 조회 및 확인할 수 있는 수량도 48매 수준에 머물렀다. "
    "둘째, 생성된 Map에 대한 불량 여부 및 유형 판정이 엔지니어의 수동 판독에 의존하여 전수 분석이 어려웠다. "
    "본 논문은 이러한 한계를 해결하는 통합 분석 아키텍처를 제안하며, 주요 기여는 다음과 같다."
)

KO_1_3 = (
    "첫째, 대량 설비 Log의 실시간 적재와 1시간 주기 Failbit Map 생성 체계를 구현하여 대량 Map 운영이 가능한 데이터 처리 기반을 구축하였다. "
    "둘째, Known 불량은 ConvNeXt V2[1] 기반 1차 wafer-level 분류와 ROI 기반 YOLO 2차 분류를 결합한 2-stage 구조로 성능을 향상시켰으며, backbone은 MaxViT[2] 대비 효율성을 기준으로 선정하였다. "
    "셋째, Unknown 불량은 SimCLR[3] 계열 self-supervised 분석과 grid-structured local sampling 및 HDBSCAN[4] 기반 grouping을 통해 보다 정밀하게 검출하도록 하였다."
)

KO_2_INTRO = "본 아키텍처는 데이터 파이프라인, Known 2-stage 분류기, Unknown self-supervised 검출기로 구성되며, 전용 분석 Web App[5]으로 배포되었다."

KO_2_1 = (
    "데이터 파이프라인은 설비 Raw Log를 실시간 적재해 1시간 주기로 Failbit Map을 생성하도록 설계하였다. "
    "핵심 병목은 wafer당 약 1,000만 개 test result의 hex-to-grade 변환 속도와 이미지 저장 용량이며, "
    "전자는 Cython으로 약 100배 향상시켰고(Fig. 1), 후자는 32-color palette-indexed PNG로 약 75% 절감하였다(Fig. 2). "
    "생성된 Map은 중앙 저장소에 적재되며 분석 Web App[5]으로 대량 조회 및 분석이 가능하다."
)

KO_2_2_A = (
    "Known 불량 분석은 16개 등록 클래스를 대상으로 하였으며, 1,500개의 Failbit Map을 사용하여 "
    "ConvNeXtV2[1] 기반 1단계 wafer-level 분류기와 저신뢰 샘플 대상 2단계 ROI(Region of Interest) 기반 YOLO를 결합한 2-stage 구조를 설계하였다."
)

KO_2_2_B = (
    "ConvNeXtV2 기반 wafer-level 분류는 전반적으로 높은 정확도와 처리 속도를 보였으나, wafer 내 불량 chip의 분포가 유사한 클래스에서는 분류 성능이 저하되었다. "
    "이를 보완하기 위해 2단계에서는 1차 분류의 저신뢰 샘플에 ROI 기반 YOLO를 적용하여 ROI 영역 내 불량 chip의 형태와 출현 패턴을 추가로 판별하도록 구성하였다(Fig. 3)."
)

KO_2_2_C = (
    "MaxViT[2]와 ConvNeXtV2 (Ref)는 동일한 weighted F1 0.87을 보였으나, ConvNeXtV2는 파라미터 수 약 26% 감소(119.5M → 88.6M)와 FLOPs 약 39% 감소(74.2G → 45.1G)로 더 효율적이라 최종 backbone으로 선정하였다. "
    "이후 Optuna를 통해 lr, weight decay, focal loss, class weight 등을 최적화하여 weighted F1을 0.92까지 향상시켰고, ROI-YOLO 보정으로 최종 0.95를 달성하였다."
)

KO_2_3_A = (
    "5일치 운영 데이터 10,000장을 사용하여 SimCLR[3] 계열 모델을 대조학습하였으며, "
    "wafer의 zone 기반 불량 해석 특성을 반영하기 위해 grid-structured local sampling을 함께 적용하였다. "
    "이후 별도 1일치 2,000장에 HDBSCAN[4]을 적용하여 유사 패턴을 그룹화하였다."
)

KO_2_3_B = (
    "Unknown 불량 검출에서는 운영 이미지 2,000장에 대한 grouping 결과 13개 후보 그룹이 검출되었으며, 현업 분석 엔지니어 검증 결과 이 중 7개가 실제 불량 그룹으로 판정되었다. "
    "나머지 6개는 일부 lot에 국한된 warning이거나 실제 불량으로 이어지지 않는 패턴이었다."
)

KO_CONCLUSION = (
    "본 연구는 Failbit Map 전수 생성과 자동 불량 분석을 위한 통합 아키텍처를 구현하였다. "
    "Cython 최적화와 32-color palette-indexed PNG로 Failbit Map의 대량 생성과 저장 및 운영을 가능하게 하였고, "
    "Known 2-stage 분류와 Unknown self-supervised 검출을 결합하여 자동 분석 흐름을 구축하였다."
)

# 본문 문단 (EN)
EN_1_1 = (
    "A Failbit Map represents defect severity at the memory-cell-block level in EDS testing using grades from 0 to 7. "
    "Because a single wafer contains roughly 10 million blocks, it is an ultra-high-resolution representation of defect location and morphology. "
    "In current practice, Measure-based analysis alone cannot capture this information, making exhaustive Failbit Map analysis necessary."
)

EN_1_2 = (
    "Two limitations have hindered practical adoption in production. "
    "First, the existing system did not have sufficient processing performance to transform equipment logs into large numbers of Failbit Maps and to store and browse them at scale. "
    "Equipment logs are typically 10 to 50 MB per wafer, and roughly 20,000 wafers are produced per day across all products. "
    "Under the previous environment, large-volume processing was limited by speed and memory constraints, and the number of maps that could be checked at once was restricted to about 48. "
    "Second, even when maps were generated, defect determination and type identification still depended on manual inspection by engineers, making exhaustive analysis difficult. "
    "To address these limitations, this paper proposes an integrated analysis architecture. The main contributions are as follows."
)

EN_1_3 = (
    "First, we implemented a real-time log ingestion and hourly Failbit Map generation pipeline, establishing a data-processing foundation for large-scale map operations. "
    "Second, for Known defects, we improved classification performance with a two-stage structure combining a ConvNeXt V2[1]-based wafer-level first-stage classifier and an ROI-based YOLO second-stage classifier; ConvNeXt V2 was selected over MaxViT[2] for efficiency. "
    "Third, for Unknown defects, we applied a SimCLR[3]-family self-supervised analysis with grid-structured local sampling and HDBSCAN[4]-based grouping for more precise detection."
)

EN_2_INTRO = "The architecture consists of a data pipeline, a two-stage Known classifier, and a self-supervised Unknown detector, deployed as a dedicated analysis web application[5]."

EN_2_1 = (
    "The data pipeline was designed to ingest equipment raw logs in real time and generate Failbit Maps on an hourly basis. "
    "The main bottlenecks were the hex-to-grade conversion speed for roughly 10 million encoded test values per wafer and the storage footprint of the resulting images; "
    "the former was accelerated by about 100x with Cython (Fig. 1), and the latter was reduced by about 75% with a 32-color palette-indexed PNG (Fig. 2). "
    "Generated maps are stored in a central repository and served through the analysis web application[5] for large-scale browsing and analysis."
)

EN_2_2_A = (
    "Known-defect analysis targeted 16 registered classes. We designed a two-stage structure combining a ConvNeXtV2[1]-based wafer-level first-stage classifier and a second-stage ROI (Region of Interest)-based YOLO for low-confidence samples, and trained it on 1,500 labeled Failbit Maps."
)

EN_2_2_B = (
    "While the ConvNeXtV2-based wafer-level classifier showed high overall accuracy and throughput, its performance degraded on classes with similar wafer-level distributions of defective chips. "
    "To compensate for this, in the second stage ROI-based YOLO re-examines low-confidence samples and identifies chip-level defect morphologies and appearance patterns within the ROI (Fig. 3)."
)

EN_2_2_C = (
    "MaxViT[2] and ConvNeXtV2 (Ref) achieved the same weighted F1 of 0.87, but ConvNeXtV2 was selected as the final backbone because it was more efficient, with about 26% fewer parameters (119.5M \u2192 88.6M) and about 39% fewer FLOPs (74.2G \u2192 45.1G). "
    "Optuna was then used to optimize fine-tuning hyperparameters such as learning rate, weight decay, focal loss, and class weights, improving weighted F1 to 0.92, "
    "and ROI-YOLO refinement further increased the final score to 0.95."
)

EN_2_3_A = (
    "Using 10,000 production maps from five days, a SimCLR[3]-family model was trained with contrastive learning using grid-structured local sampling to reflect the zone-based nature of wafer defect interpretation. "
    "HDBSCAN[4] was then applied to a separate set of 2,000 maps from one additional day to group similar patterns."
)

EN_2_3_B = (
    "For Unknown defects, grouping on 2,000 production images yielded 13 candidate groups, and validation by domain engineers confirmed that 7 of them were real defect groups. "
    "The remaining 6 were lot-specific warnings or non-defect patterns."
)

EN_CONCLUSION = (
    "This study implemented an integrated architecture for end-to-end Failbit Map generation and automated defect analysis. "
    "Cython optimization and a 32-color palette-indexed PNG enabled large-scale map generation, storage, and operational use, "
    "while the combination of Known two-stage classification and Unknown self-supervised discovery established an automated defect-analysis workflow."
)

FIG3_CAP = (
    "Fig. 3. Representative patterns of Class A(a) and Class B(b), and a true Class A sample(c) "
    "misclassified as Class B by the first-stage CNN but corrected to Class A by the second-stage ROI-YOLO. "
    "The blue circle indicates the ROI, and the red boxes denote YOLO detections of chip-level defect patterns within the ROI."
)
FIG4_CAP = "Fig. 4. Unknown-fail grouping on production images."


def _setup_doc(default_font: str) -> Document:
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(21)
    sec.page_height = Cm(29.7)
    set_margins(sec)
    set_two_columns(sec, space_cm=0.5)
    doc.styles["Normal"].font.size = PT10
    doc.styles["Normal"].font.name = default_font

    main_sp = sec._sectPr
    tp = OxmlElement("w:type")
    tp.set(qn("w:val"), "continuous")
    main_sp.insert(0, tp)
    return doc


def _build(lang: str) -> Document:
    is_ko = (lang == "ko")
    doc = _setup_doc(F_KO if is_ko else F_EN)

    if is_ko:
        add_title_block(doc, TITLE_KO_CODEX, TITLE_EN_CODEX)
    else:
        # 영문본: 영어 제목만 크게
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(TITLE_EN_CODEX)
        from generate_papers import _apply_run_font, _set_single_spacing, PT20
        _apply_run_font(r, size=PT20, bold=True)
        _set_single_spacing(p)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(8)

    add_author_block(doc, AUTHORS if is_ko else AUTHORS_EN,
                     AFFIL if is_ko else AFFIL_EN)
    add_abstract_block(doc, ABSTRACT_KO if is_ko else ABSTRACT_EN,
                       label="(Abstract)")

    from generate_papers import _apply_run_font, _set_single_spacing
    pk = doc.add_paragraph()
    rk = pk.add_run(KEYWORDS_LINE)
    _apply_run_font(rk, size=PT9, italic=True)
    _set_single_spacing(pk)
    pk.paragraph_format.space_after = Pt(4)

    add_1col_header_break(doc)

    add_heading(doc, "1. INTRODUCTION", level=1)
    add_body(doc, KO_1_1 if is_ko else EN_1_1, indent=True, space_after=Pt(2))
    add_body(doc, KO_1_2 if is_ko else EN_1_2, indent=True, space_after=Pt(2))
    add_body(doc, KO_1_3 if is_ko else EN_1_3, indent=True, space_after=Pt(2))

    add_heading(doc, "2. PROPOSED METHOD", level=1)
    add_body(doc, KO_2_INTRO if is_ko else EN_2_INTRO, indent=True, space_after=Pt(2))

    add_subheading(doc, "2.1 DATA PIPELINE")
    add_body(doc, KO_2_1 if is_ko else EN_2_1, space_after=Pt(2))

    add_labeled_example_table(
        doc,
        "Hex-to-grade conversion",
        [
            ("Raw:", "090B0C0D0E0F090A0B0C"),
            ("Decoding:", "\"0C\" -> \"C\" -> 12 (hex to decimal) -> 3"),
            ("Python:", "interpreter-based loop execution"),
            ("Cython:", "compiled integer loop execution"),
            ("Grade:", "0 2 3 4 5 6 0 1 2 3"),
        ],
        "Fig. 1. Hex-to-grade conversion accelerated by Cython",
    )
    add_labeled_example_table(
        doc,
        "RGB PNG vs Palette-indexed PNG",
        [
            ("RGB PNG:", [
                "[(123,54,24), (123,54,24), ..., (123,54,24)],",
                "[(123,54,24), (123,54,24), ..., (123,54,24)]",
            ]),
            ("Palette-indexed PNG:", [
                "P[3] = (123,54,24)",
                "[(3), (3), ..., (3)],",
                "[(3), (3), ..., (3)]",
            ]),
            ("RGB_to_Palette:", "(123,54,24) -> (3)"),
        ],
        "Fig. 2. Palette-indexed PNG for failbit map compression.",
    )

    add_subheading(doc, "2.2 Known 불량 분류" if is_ko else "2.2 Known Defect Classification")
    add_body(doc, KO_2_2_A if is_ko else EN_2_2_A, space_after=Pt(2))
    add_body(doc, KO_2_2_B if is_ko else EN_2_2_B, space_after=Pt(2))

    _fig_known_path = str(OUT_DIR / "_fig_yolo_roi.png")
    if not Path(_fig_known_path).exists():
        make_yolo_roi_figure(_fig_known_path)
    add_figure(doc, _fig_known_path, FIG3_CAP, width_cm=8.2)

    # Table 1 — 영문 캡션에서도 weighted F1 소문자 통일
    table_perf = dict(TABLE_PERF_CLAUDE)
    table_perf["caption"] = (
        "Table 1. Backbone comparison and staged improvements for "
        "known-fail classification (16-class, test weighted F1)"
    )
    add_table(doc, table_perf)

    add_body(doc, KO_2_2_C if is_ko else EN_2_2_C, space_after=Pt(2))

    add_subheading(doc, "2.3 Unknown 불량 검출" if is_ko else "2.3 Unknown Defect Discovery")
    add_body(doc, KO_2_3_A if is_ko else EN_2_3_A, space_after=Pt(2))

    _fig_unknown_path = str(OUT_DIR / "_fig_cluster.png")
    add_figure(doc, _fig_unknown_path, FIG4_CAP, width_cm=8.2)

    add_body(doc, KO_2_3_B if is_ko else EN_2_3_B, space_after=Pt(2))

    add_heading(doc, "3. CONCLUSION", level=1)
    add_body(doc, KO_CONCLUSION if is_ko else EN_CONCLUSION, indent=True, space_after=Pt(2))

    add_refs(doc, REFS_CODEX, title="REFERENCES")
    return doc


if __name__ == "__main__":
    doc_ko = _build("ko")
    out_ko = OUT_DIR / "paper_codex_2page_rev167.docx"
    doc_ko.save(str(out_ko))
    print(f"saved: {out_ko}")

    doc_en = _build("en")
    out_en = OUT_DIR / "paper_codex_2page_rev167_en.docx"
    doc_en.save(str(out_en))
    print(f"saved: {out_en}")
