from __future__ import annotations

from pathlib import Path
from typing import Iterable, List

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[1]
SOURCE_MD = ROOT / "paper_submission_2page.md"
OUTPUT_DOCX = ROOT / "paper_submission_2page.docx"


def set_run_font(run, size_pt: int, bold: bool = False) -> None:
    run.bold = bold
    run.font.size = Pt(size_pt)
    run.font.name = "Times New Roman"
    rfonts = run._element.rPr.rFonts
    rfonts.set(qn("w:ascii"), "Times New Roman")
    rfonts.set(qn("w:hAnsi"), "Times New Roman")
    rfonts.set(qn("w:eastAsia"), "바탕체")


def set_paragraph_single_spacing(paragraph) -> None:
    pf = paragraph.paragraph_format
    pf.line_spacing = 1.0
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)


def add_text_paragraph(
    doc_or_cell,
    text: str,
    *,
    size: int = 10,
    bold: bool = False,
    align=WD_ALIGN_PARAGRAPH.LEFT,
    before_pt: int = 0,
    after_pt: int = 0,
) -> None:
    p = doc_or_cell.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.line_spacing = 1.0
    pf.space_before = Pt(before_pt)
    pf.space_after = Pt(after_pt)
    run = p.add_run(text)
    set_run_font(run, size, bold=bold)


def set_page_layout(section) -> None:
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(3.0)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)
    section.gutter = Cm(0)


def set_columns(section, count: int, space_cm: float) -> None:
    sect_pr = section._sectPr
    cols = sect_pr.xpath("./w:cols")
    cols_el = cols[0] if cols else OxmlElement("w:cols")
    cols_el.set(qn("w:num"), str(count))
    cols_el.set(qn("w:space"), str(int(space_cm * 567)))
    if not cols:
        sect_pr.append(cols_el)


def read_lines() -> List[str]:
    return SOURCE_MD.read_text(encoding="utf-8").splitlines()


def extract_title(lines: List[str]) -> str:
    for line in lines:
        if line.startswith("# "):
            return line[2:].strip()
    raise ValueError("Title not found")


def section_lines(lines: List[str], heading: str) -> List[str]:
    start = None
    for idx, line in enumerate(lines):
        if line.strip() == heading:
            start = idx + 1
            break
    if start is None:
        return []
    end = len(lines)
    for idx in range(start, len(lines)):
        if lines[idx].startswith("## ") and lines[idx].strip() != heading:
            end = idx
            break
    return lines[start:end]


def parse_blocks(lines: Iterable[str]) -> List[dict]:
    items = list(lines)
    blocks: List[dict] = []
    i = 0
    while i < len(items):
        raw = items[i]
        line = raw.strip()
        if not line or line == "---":
            i += 1
            continue
        if line.startswith("### "):
            blocks.append({"type": "heading", "level": 3, "text": line[4:].strip()})
            i += 1
            continue
        if line.startswith("## "):
            blocks.append({"type": "heading", "level": 2, "text": line[3:].strip()})
            i += 1
            continue
        if line.startswith("**Table ") and line.endswith("**"):
            blocks.append({"type": "caption", "kind": "table", "text": line.strip("*")})
            i += 1
            continue
        if line.startswith("**Table "):
            blocks.append({"type": "caption", "kind": "table", "text": line.replace("**", "")})
            i += 1
            continue
        if line.startswith("**Figure "):
            blocks.append({"type": "caption", "kind": "figure", "text": line.replace("**", "")})
            i += 1
            continue
        if line.startswith("|"):
            table_lines = []
            while i < len(items) and items[i].strip().startswith("|"):
                table_lines.append(items[i].strip())
                i += 1
            blocks.append({"type": "table", "lines": table_lines})
            continue
        if line.startswith("- "):
            bullet_lines = []
            while i < len(items) and items[i].strip().startswith("- "):
                bullet_lines.append(items[i].strip()[2:].strip())
                i += 1
            blocks.append({"type": "bullets", "items": bullet_lines})
            continue
        if line.startswith("[") and "]" in line[:5]:
            blocks.append({"type": "reference", "text": line})
            i += 1
            continue
        if line.startswith("**Keywords:**"):
            blocks.append({"type": "keywords", "text": line.replace("**", "")})
            i += 1
            continue

        para_lines = [line]
        i += 1
        while i < len(items):
            nxt = items[i].strip()
            if not nxt or nxt == "---" or nxt.startswith(("## ", "### ", "- ", "|", "**Table ", "**Figure ", "[", "**Keywords:**")):
                break
            para_lines.append(nxt)
            i += 1
        blocks.append({"type": "paragraph", "text": " ".join(para_lines)})
    return blocks


def add_markdown_table(doc: Document, lines: List[str]) -> None:
    rows = []
    for idx, line in enumerate(lines):
        parts = [p.strip() for p in line.strip("|").split("|")]
        if idx == 1 and all(set(p) <= {"-", ":"} for p in parts):
            continue
        rows.append(parts)
    if not rows:
        return
    col_count = len(rows[0])
    table = doc.add_table(rows=0, cols=col_count)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    widths = [Cm(2.8), Cm(4.2), Cm(3.0), Cm(2.2)]
    if col_count != 4:
        widths = [Cm((17.5 / col_count)) for _ in range(col_count)]
    for r_idx, row_data in enumerate(rows):
        row = table.add_row()
        for c_idx, cell_text in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.width = widths[c_idx]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            set_paragraph_single_spacing(p)
            run = p.add_run(cell_text)
            set_run_font(run, 9, bold=(r_idx == 0))


def add_bullets(doc: Document, items: List[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        pf = p.paragraph_format
        pf.line_spacing = 1.0
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        run = p.add_run(item)
        set_run_font(run, 10, bold=False)


def build_doc() -> Document:
    lines = read_lines()
    title = extract_title(lines)
    abstract_blocks = parse_blocks(section_lines(lines, "## 초록"))
    body_blocks = [{"type": "heading", "level": 2, "text": "1. 서론"}]
    body_blocks += parse_blocks(section_lines(lines, "## 1. 서론"))
    body_blocks += [{"type": "heading", "level": 2, "text": "2. 제안 방법"}]
    body_blocks += parse_blocks(section_lines(lines, "## 2. 제안 방법"))
    body_blocks += [{"type": "heading", "level": 2, "text": "3. 실험 결과 및 논의"}]
    body_blocks += parse_blocks(section_lines(lines, "## 3. 실험 결과 및 논의"))
    body_blocks += [{"type": "heading", "level": 2, "text": "4. 결론"}]
    body_blocks += parse_blocks(section_lines(lines, "## 4. 결론"))
    body_blocks += [{"type": "heading", "level": 2, "text": "참고문헌"}]
    body_blocks += parse_blocks(section_lines(lines, "## 참고문헌"))

    doc = Document()
    set_page_layout(doc.sections[0])

    add_text_paragraph(doc, title, size=20, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after_pt=6)
    add_text_paragraph(
        doc,
        "Draft document: example performance values should be replaced with actual server logs before submission.",
        size=9,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        after_pt=8,
    )
    add_text_paragraph(doc, "Abstract", size=10, bold=True, before_pt=6, after_pt=6)

    for block in abstract_blocks:
        if block["type"] == "paragraph":
            add_text_paragraph(doc, block["text"], size=10)
        elif block["type"] == "keywords":
            add_text_paragraph(doc, block["text"], size=10, before_pt=6)

    body_section = doc.add_section(WD_SECTION_START.CONTINUOUS)
    set_page_layout(body_section)
    set_columns(body_section, count=2, space_cm=0.5)

    for block in body_blocks:
        btype = block["type"]
        if btype == "heading":
            level = block["level"]
            text = block["text"]
            size = 11
            add_text_paragraph(doc, text, size=size, bold=True, before_pt=10, after_pt=10)
        elif btype == "paragraph":
            add_text_paragraph(doc, block["text"], size=10)
        elif btype == "bullets":
            add_bullets(doc, block["items"])
        elif btype == "caption":
            add_text_paragraph(doc, block["text"], size=9, bold=True, before_pt=10, after_pt=6)
        elif btype == "table":
            add_markdown_table(doc, block["lines"])
            add_text_paragraph(doc, "", size=10, after_pt=10)
        elif btype == "reference":
            add_text_paragraph(doc, block["text"], size=9)

    return doc


def main() -> None:
    doc = build_doc()
    doc.save(OUTPUT_DOCX)
    print(f"Saved {OUTPUT_DOCX}")


if __name__ == "__main__":
    main()
