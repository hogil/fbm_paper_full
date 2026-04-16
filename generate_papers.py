#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
논문 Word 파일 생성 스크립트
- paper_detailed.docx : 상세 버전
- paper_2page.docx    : 2페이지 압축 버전

레이아웃 규칙:
  제목/저자/초록 : 1단 (전체 폭)
  본문~참고문헌  : 2단 (0.5cm 간격)
  여백: 위 3.0 / 아래 2.5 / 좌우 1.5 cm
  서체: 바탕체 (국문) / Times New Roman (영문)
  줄간격: 1.0 (single)
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT_DIR = Path(__file__).parent

# ─── 폰트 & 크기 ─────────────────────────────────────────────
F_KO  = "바탕체"
F_EN  = "Times New Roman"

PT20  = Pt(20)   # 제목
PT12  = Pt(12)   # 저자명
PT11  = Pt(11)   # 본문 제목 (섹션)
PT10  = Pt(10)   # 본문 내용 / 초록
PT9   = Pt(9)    # 그림·표 캡션 / 참고문헌
PT8   = Pt(8)    # 수식/예시 블록 소형

# ─── 여백 (cm) ───────────────────────────────────────────────
M_TOP    = 3.0
M_BOTTOM = 2.5
M_LEFT   = 1.5
M_RIGHT  = 1.5

# ─── 문서별 데이터 ────────────────────────────────────────────
TITLE_KO_CLAUDE = "DRAM Failbit Map 기반 Known·Unknown 불량 분류·검출 아키텍처"
TITLE_EN_CLAUDE = ("A Known/Unknown Defect Classification and Detection Architecture "
                   "Based on DRAM Failbit Maps")

TITLE_KO_CODEX = "DRAM Failbit Map 기반 Known 불량 및 Unknown 불량 분석 AI 아키텍처"
TITLE_EN_CODEX = ("An AI Architecture for Known and Unknown Defect Analysis "
                  "Based on DRAM Failbit Maps")

AUTHORS  = "홍길동\u00b9, 김철수\u00b9"   # 위첨자 번호 포함 — 실제 이름으로 교체
AFFIL    = "\u00b9 반도체연구소, Samsung Electronics, 화성시, 대한민국"

ABSTRACT_CLAUDE = (
    "Failbit Map은 반도체 EDS Test에서 생성되는 웨이퍼당 약 1,000만 pixel 규모의 초고해상도 공간 데이터로, "
    "불량 분석의 핵심 자료이다. 그러나 현업에서는 대량 Map의 생성·조회 환경이 부족하고 "
    "불량 판단이 엔지니어의 수작업에 의존해 전수 분석이 어렵다. "
    "본 논문은 대규모 Failbit Map 생성·저장 파이프라인과 "
    "Known·Unknown 불량 통합 분석 AI 아키텍처를 End-to-End 구조로 구현하였다. "
    "Known 경로는 ConvNeXtV2-Base와 저신뢰 샘플 선택적 ROI-YOLO를 결합한 2단계 구조로 "
    "16-class weighted F1 0.78→0.95를 달성하였다. "
    "Unknown 경로는 SimCLR 계열 contrastive learning과 grid structured local sampling으로 "
    "5일치 10,000장을 학습한 뒤 1일치 2,000장을 13개 후보 그룹으로 압축하였으며(154배 축소), "
    "현업 엔지니어 검증을 통해 7개가 실제 불량 그룹으로 확인되었다."
)

ABSTRACT_CODEX = (
    "Failbit Map은 반도체 EDS Test에서 생성되는 웨이퍼당 약 1,000만 pixel 수준의 초고해상도 데이터로, 불량 패턴 분석의 핵심 자료이다. "
    "그러나 실제 현업에서는 대량의 Failbit Map 조회가 불가능하고, 일부 Map 분석도 엔지니어의 수작업에 의존하고 있다. "
    "본 논문은 이를 해결하기 위해 대규모 Failbit Map을 생성 및 저장하는 파이프라인과 Known 불량 분류와 Unknown 불량 검출 AI 아키텍처를 구현하였다. "
    "Cython 적용으로 데이터 변환 속도를 약 100배 향상시켰고, Palette PNG 적용으로 이미지 용량을 약 75% 절감하였다. "
    "Known 불량 분류는 ConvNeXtV2 기반 1차 분류와 저신뢰 샘플에 대한 ROI 기반 YOLO 2차 분류를 결합한 구조로 설계하였으며, F1-score 0.95를 달성하였다. "
    "Unknown 불량 검출은 레이블 없이 SimCLR 계열 contrastive learning 기반으로 수행하였고, wafer의 zone 기반 불량 해석 특성을 반영하기 위해 grid structured local sampling을 적용하였다. "
    "실제 양산 5일치 Failbit Map 10,000장을 학습한 뒤 1일치 2,000장에 대해 추론한 결과 13개 불량 그룹이 검출되었으며, 현업 분석 엔지니어의 검증 결과 이 중 7개가 실제 불량 그룹으로 판정되어, 실제 운영 환경에서의 적용 가능성을 입증하였다."
)

# ─── 표 데이터 ───────────────────────────────────────────────

TABLE1 = {
    "caption": "Table 1. 등록 불량 성능 향상 단계 (Weighted F1)",
    "headers": ["단계", "구성", "Weighted F1"],
    "rows": [
        ["Baseline CNN",              "초기 기준 모델",                           "0.78"],
        ["ConvNeXtV2-Base (HF timm)", "backbone 교체, 384×384, IN-22k 사전학습", "0.85"],
        ["+ Optuna 하이퍼파라미터 탐색", "LR·WD·scheduler·augmentation 최적화",  "0.92"],
        ["+ ROI enhancement + YOLO",  "선택적 2-stage correction",               "0.95"],
    ],
}

TABLE_BACKBONE = {
    "caption": "Table 2. Backbone 모델 비교 (Hugging Face timm fine-tuning, Val Weighted F1)",
    "headers": ["모델", "사전학습", "Input", "Val F1", "비고"],
    "rows": [
        ["ViT-B/16",          "IN-21k",      "224", "0.80", "소규모 데이터에서 attention 학습 부족"],
        ["EfficientNetV2-M",  "IN-1k",       "224", "0.82", "경량 구조, 공간 패턴 표현 한계"],
        ["Swin-B",            "IN-1k",       "224", "0.83", "shifted window로 지역성 보완"],
        ["MaxViT-T",          "IN-1k",       "224", "0.84", "multi-axis attention"],
        ["ConvNeXtV2-Base",   "FCMAE+IN-22k","384", "0.92", "1위, FCMAE 비지도 사전학습 강건"],
    ],
}

# claude 버전 전용 — codex rev9 기준 (384 입력, 단순화 컬럼)
TABLE_BACKBONE_CLAUDE = {
    "caption": "Table 1. Backbone 모델 비교 (test Weighted F1, input 384×384)",
    "headers": ["Backbone", "Pretraining", "Test F1"],
    "rows": [
        ["ViT-Base/16",       "IN-21k",          "0.81"],
        ["Swin-Base",         "IN-1k",           "0.84"],
        ["EfficientNetV2-M",  "IN-1k",           "0.85"],
        ["MaxViT-Base",       "IN-1k",           "0.87"],
        ["ConvNeXtV2-Base",   "FCMAE + IN-22k",  "0.92"],
    ],
}

TABLE1_CLAUDE = {
    "caption": "Table 2. 등록 불량 분류 단계별 test Weighted F1",
    "headers": ["Stage", "Configuration", "Test F1"],
    "rows": [
        ["Baseline",  "ResNet-50 (224×224)",              "0.78"],
        ["Stage 1",   "ConvNeXtV2-Base (backbone only)",  "0.85"],
        ["Stage 2",   "+ Optuna HPO",                     "0.92"],
        ["Stage 3",   "+ Grad-CAM ROI + YOLO",            "0.95"],
    ],
}

# 2-page claude 전용 — 단순 3행 성능 요약
TABLE_PERF_CLAUDE = {
    "caption": "Table 1. Backbone comparison and staged improvements for known-defect classification (16-class, test Weighted F1)",
    "headers": ["구성", "사전학습", "Test F1", "비고"],
    "font_size": 8,
    "widths_cm": [3.3, 1.9, 1.0, 1.6],
    "rows": [
        ["Baseline CNN",         "-",            "0.78", "baseline"],
        ["ViT",                  "IN-21k",       "0.81", "fine-tune"],
        ["Swin",                 "IN-1k",        "0.84", "fine-tune"],
        ["EffNetV2",             "IN-1k",        "0.85", "fine-tune"],
        ["MaxViT",               "IN-21k",       "0.87", "fine-tune"],
        ["ConvNeXtV2 (Ref)",     "IN-22k",       "0.87", "selected"],
        ["Ref + Optuna",         "IN-22k",       "0.92", "HPO"],
        ["Ref + Optuna + ROI",   "IN-22k",       "0.95", "2-stage"],
    ],
}

TABLE2A = {
    "caption": "Table 3. Confusion Matrix 개선 상위 클래스 쌍 (1→2단계)",
    "headers": ["Stage1 예측", "실제(True)", "교정 건수", "교정률", "ROI 역할"],
    "rows": [
        ["arc",       "ring",     "35건", "71%", "외곽 호 형태 집중"],
        ["edge_loc",  "edge_ring","47건", "68%", "외곽 경계 집중"],
        ["edge_ring", "edge_loc", "39건", "63%", "위치 패턴 구분"],
        ["local",     "random",   "31건", "55%", "군집 위치 확인"],
        ["random",    "local",    "28건", "51%", "분산 패턴 구분"],
        ["h_line",    "v_line",   "19건", "52%", "방향성 식별"],
        ["near_full", "donut",    "22건", "44%", "중심 공백 확인"],
    ],
}

TABLE2 = {
    "caption": "Table 4. 클래스별 Weighted F1 (16-class, 2단계 파이프라인)",
    "headers": ["클래스", "F1 1단계", "F1 2단계", "Δ"],
    "rows": [
        ["arc",            "0.83", "0.89", "+0.06"],
        ["center",         "0.95", "0.97", "+0.02"],
        ["cluster",        "0.88", "0.92", "+0.04"],
        ["donut",          "0.93", "0.95", "+0.02"],
        ["edge_loc",       "0.87", "0.92", "+0.05"],
        ["edge_ring",      "0.86", "0.91", "+0.05"],
        ["h_line",         "0.87", "0.91", "+0.04"],
        ["line",           "0.91", "0.94", "+0.03"],
        ["local",          "0.85", "0.90", "+0.05"],
        ["near_full",      "0.94", "0.96", "+0.02"],
        ["none",           "0.97", "0.98", "+0.01"],
        ["random",         "0.91", "0.93", "+0.02"],
        ["ring",           "0.89", "0.93", "+0.04"],
        ["scratch",        "0.96", "0.97", "+0.01"],
        ["spot",           "0.90", "0.93", "+0.03"],
        ["v_line",         "0.86", "0.90", "+0.04"],
        ["Weighted avg",   "0.92", "0.95", "+0.03"],
    ],
}

TABLE_THRESH = {
    "caption": "Table 5. threshold 및 YOLO 설정",
    "headers": ["항목", "실험 범위", "채택값"],
    "rows": [
        ["Classifier confidence",     "0.70 ~ 0.90", "0.80"],
        ["Difficult class threshold", "0.75 ~ 0.85", "0.80"],
        ["YOLO confidence",           "0.15 ~ 0.35", "0.25"],
        ["YOLO cls gain",             "0.5 ~ 3.0",   "1.5"],
        ["YOLO imgsz",                "512 ~ 768",   "640"],
        ["YOLO batch",                "8 ~ 32",       "16"],
        ["YOLO epochs",               "50 ~ 150",    "100"],
    ],
}

TABLE_COMPARE = {
    "caption": "Table 6. 방법론 비교",
    "headers": ["항목", "ResNet-50", "CNN+K-means", "본 연구"],
    "rows": [
        ["등록 불량 분류",            "O",    "O",    "O"],
        ["미등록 불량 탐지",          "X",    "△",    "O"],
        ["Measure(FTN·QTN) 분석 연동","X",    "X",    "O"],
        ["클러스터 수 자동",          "—",    "X",    "O"],
        ["설명 가능성(Grad-CAM)",     "X",    "X",    "O"],
        ["Test Weighted F1",          "0.78", "0.86", "0.95"],
    ],
}

TABLE_HDBSCAN = {
    "caption": "Table 7. HDBSCAN 주요 파라미터 구성 비교",
    "headers": ["구성", "min_cluster_size", "epsilon", "군집 수", "무시 비율"],
    "rows": [
        ["세분화 (aggressive)",  "5",  "0.02", "多",  "낮음"],
        ["채택 (기본)",          "12", "0.06", "적정","적정"],
        ["보수적 (conservative)","25", "0.15", "少",  "높음"],
    ],
}

REFS = [
    "[1] T. Nakazawa and D. V. Kulkarni, \"Wafer Map Defect Pattern Classification and Image Retrieval "
    "Using Convolutional Neural Network,\" IEEE Transactions on Semiconductor Manufacturing, vol. 31, "
    "no. 2, pp. 309-314, 2018.",
    "[2] M. B. Alawieh, D. Boning, and D. Z. Pan, \"Wafer Map Defect Patterns Classification Using Deep "
    "Selective Learning,\" in Proceedings of the 57th ACM/IEEE Design Automation Conference (DAC), "
    "pp. 1-6, 2020.",
    "[3] J. Jang and G. T. Lee, \"Decision fusion approach for detecting unknown wafer bin map patterns "
    "based on a deep multitask learning model,\" Expert Systems with Applications, vol. 215, "
    "art. 119363, 2023.",
    "[4] S. Woo et al., \"ConvNeXt V2: Co-Designing and Scaling ConvNets With Masked Autoencoders,\" "
    "in Proceedings of the IEEE/CVF Conference on Computer Vision and Pattern Recognition (CVPR), "
    "pp. 16133-16142, 2023.",
    "[5] R. R. Selvaraju et al., \"Grad-CAM: Visual Explanations From Deep Networks via Gradient-Based "
    "Localization,\" in Proceedings of the IEEE International Conference on Computer Vision (ICCV), "
    "pp. 618-626, 2017.",
    "[6] T. Chen, S. Kornblith, M. Norouzi, and G. Hinton, \"A Simple Framework for Contrastive Learning "
    "of Visual Representations,\" in Proceedings of the 37th International Conference on Machine Learning "
    "(ICML), PMLR 119, pp. 1597-1607, 2020.",
    "[7] R. J. G. B. Campello, D. Moulavi, and J. Sander, \"Density-Based Clustering Based on Hierarchical "
    "Density Estimates,\" in Advances in Knowledge Discovery and Data Mining, PAKDD 2013, pp. 160-172, 2013.",
    "[8] Hugging Face, \"facebook/convnextv2-base-22k-384,\" model card, accessed Apr. 15, 2026. "
    "Available: https://huggingface.co/facebook/convnextv2-base-22k-384",
    "[9] Hugging Face timm, \"maxvit_base_tf_384.in21k_ft_in1k,\" model card, accessed Apr. 15, 2026. "
    "Available: https://huggingface.co/timm/maxvit_base_tf_384.in21k_ft_in1k",
    "[10] Z. Tu et al., \"MaxViT: Multi-Axis Vision Transformer,\" in Proceedings of the European "
    "Conference on Computer Vision (ECCV), pp. 459-479, 2022.",
]

REFS_CODEX = [
    "[1] T. Nakazawa and D. V. Kulkarni, \"Wafer Map Defect Pattern Classification and Image Retrieval "
    "Using Convolutional Neural Network,\" IEEE Transactions on Semiconductor Manufacturing, vol. 31, "
    "no. 2, pp. 309-314, 2018.",
    "[2] M. B. Alawieh, D. Boning, and D. Z. Pan, \"Wafer Map Defect Patterns Classification Using Deep "
    "Selective Learning,\" in Proceedings of the 57th ACM/IEEE Design Automation Conference (DAC), "
    "pp. 1-6, 2020.",
    "[3] J. Jang and G. T. Lee, \"Decision fusion approach for detecting unknown wafer bin map patterns "
    "based on a deep multitask learning model,\" Expert Systems with Applications, vol. 215, "
    "art. 119363, 2023.",
    "[4] S. Woo et al., \"ConvNeXt V2: Co-Designing and Scaling ConvNets With Masked Autoencoders,\" "
    "in Proceedings of the IEEE/CVF Conference on Computer Vision and Pattern Recognition (CVPR), "
    "pp. 16133-16142, 2023.",
    "[5] T. Chen, S. Kornblith, M. Norouzi, and G. Hinton, \"A Simple Framework for Contrastive Learning "
    "of Visual Representations,\" in Proceedings of the 37th International Conference on Machine Learning "
    "(ICML), PMLR 119, pp. 1597-1607, 2020.",
    "[6] R. J. G. B. Campello, D. Moulavi, and J. Sander, \"Density-Based Clustering Based on Hierarchical "
    "Density Estimates,\" in Advances in Knowledge Discovery and Data Mining, PAKDD 2013, pp. 160-172, 2013.",
    "[7] Z. Tu et al., \"MaxViT: Multi-Axis Vision Transformer,\" in Proceedings of the European "
    "Conference on Computer Vision (ECCV), pp. 459-479, 2022.",
]

# build_2page_claude 전용 참고문헌
# [3] SupCon(미사용) → Optuna 교체, [9] YOLO 추가
REFS_CLAUDE = [
    "[1] T. Nakazawa and D. V. Kulkarni, \"Wafer Map Defect Pattern Classification and Image Retrieval "
    "Using Convolutional Neural Network,\" IEEE Trans. Semiconductor Manufacturing, vol. 31, no. 2, "
    "pp. 309-314, 2018.",
    "[2] M. B. Alawieh, D. Boning, and D. Z. Pan, \"Wafer Map Defect Patterns Classification Using Deep "
    "Selective Learning,\" in Proc. 57th ACM/IEEE DAC, pp. 1-6, 2020.",
    "[3] T. Akiba et al., \"Optuna: A Next-generation Hyperparameter Optimization Framework,\" "
    "in Proc. 25th ACM SIGKDD, pp. 2623-2631, 2019.",
    "[4] J. Jang and G. T. Lee, \"Decision Fusion Approach for Detecting Unknown Wafer Bin Map Patterns "
    "Based on a Deep Multitask Learning Model,\" Expert Systems with Applications, vol. 215, art. 119363, 2023.",
    "[5] S. Woo et al., \"ConvNeXt V2: Co-Designing and Scaling ConvNets With Masked Autoencoders,\" "
    "in Proc. CVPR, pp. 16133-16142, 2023.",
    "[6] R. R. Selvaraju et al., \"Grad-CAM: Visual Explanations From Deep Networks via Gradient-Based "
    "Localization,\" in Proc. ICCV, pp. 618-626, 2017.",
    "[7] T. Chen et al., \"A Simple Framework for Contrastive Learning of Visual Representations,\" "
    "in Proc. ICML, PMLR 119, pp. 1597-1607, 2020.",
    "[8] R. J. G. B. Campello et al., \"Density-Based Clustering Based on Hierarchical Density Estimates,\" "
    "in PAKDD 2013, pp. 160-172, 2013.",
    "[9] J. Redmon et al., \"You Only Look Once: Unified, Real-Time Object Detection,\" "
    "in Proc. CVPR, pp. 779-788, 2016.",
]


# ══════════════════════════════════════════════════════════════
# 헬퍼 함수
# ══════════════════════════════════════════════════════════════

def _apply_run_font(run, size=PT10, bold=False, italic=False):
    run.font.name  = F_KO
    run.font.size  = size
    run.font.bold  = bold
    run.font.italic = italic
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), F_KO)
    rFonts.set(qn("w:ascii"),    F_EN)
    rFonts.set(qn("w:hAnsi"),    F_EN)


def _set_single_spacing(p):
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE


def set_margins(section, top=M_TOP, bottom=M_BOTTOM, left=M_LEFT, right=M_RIGHT):
    section.top_margin    = Cm(top)
    section.bottom_margin = Cm(bottom)
    section.left_margin   = Cm(left)
    section.right_margin  = Cm(right)


def set_two_columns(section, space_cm: float = 0.5):
    sectPr = section._sectPr
    cols_elem = None
    for child in sectPr:
        if child.tag == qn("w:cols"):
            cols_elem = child
            break
    if cols_elem is None:
        cols_elem = OxmlElement("w:cols")
        sectPr.append(cols_elem)
    twips = int(space_cm * 567)
    cols_elem.set(qn("w:num"),        "2")
    cols_elem.set(qn("w:space"),      str(twips))
    cols_elem.set(qn("w:equalWidth"), "1")


def add_single_to_double_break(doc):
    """
    [기존 함수 — build_detailed / build_2page 에서 사용, 변경 금지]
    inline sectPr 에 type=continuous 포함 (구버전 방식).
    """
    from copy import deepcopy

    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after  = Pt(0)
    r = p.add_run()
    r.font.size = Pt(1)

    pPr = p._p.get_or_add_pPr()
    sectPr = OxmlElement("w:sectPr")

    tp = OxmlElement("w:type")
    tp.set(qn("w:val"), "continuous")
    sectPr.append(tp)

    main_sp = doc.sections[0]._sectPr
    for tag in ("w:pgSz", "w:pgMar", "w:docGrid"):
        el = main_sp.find(qn(tag))
        if el is not None:
            sectPr.append(deepcopy(el))

    cols = OxmlElement("w:cols")
    cols.set(qn("w:num"), "1")
    sectPr.append(cols)

    pPr.append(sectPr)
    return p


def add_1col_header_break(doc):
    """
    Word 직접 생성 XML 분석으로 확인된 올바른 1단→2단 연속 전환.

    핵심 발견:
      - inline sectPr (1단 섹션 끝): type 없음, cols=1 만
      - 최종 sectPr (2단 본문): type=continuous ← 여기에 있어야 함
    type=continuous 를 inline sectPr 에 넣으면 Word 가 강제 page break 함.
    """
    from copy import deepcopy

    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after  = Pt(0)
    r = p.add_run()
    r.font.size = Pt(1)

    pPr = p._p.get_or_add_pPr()
    sectPr = OxmlElement("w:sectPr")

    # type 없음 — 1단 섹션은 첫 번째 섹션이므로 type 불필요
    # pgSz / pgMar / docGrid 는 최종 sectPr 에서 복사 (twip 불일치 방지)
    main_sp = doc.sections[0]._sectPr
    for tag in ("w:pgSz", "w:pgMar", "w:docGrid"):
        el = main_sp.find(qn(tag))
        if el is not None:
            sectPr.append(deepcopy(el))

    cols = OxmlElement("w:cols")
    cols.set(qn("w:num"), "1")
    sectPr.append(cols)

    pPr.append(sectPr)
    return p


# ── 콘텐츠 추가 함수 ─────────────────────────────────────────

def add_title_block(doc, ko: str, en: str):
    """제목 20pt Bold 중앙 정렬"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(ko)
    _apply_run_font(r, size=PT20, bold=True)
    _set_single_spacing(p)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(4)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(en)
    _apply_run_font(r2, size=PT10, italic=True)
    _set_single_spacing(p2)
    p2.paragraph_format.space_after = Pt(4)


def add_author_block(doc, authors: str, affil: str):
    """저자명 12pt Bold / 소속 10pt — 중앙 정렬"""
    pa = doc.add_paragraph()
    pa.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ra = pa.add_run(authors)
    _apply_run_font(ra, size=PT12, bold=True)
    _set_single_spacing(pa)
    pa.paragraph_format.space_after = Pt(2)

    pf = doc.add_paragraph()
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rf = pf.add_run(affil)
    _apply_run_font(rf, size=PT9)
    _set_single_spacing(pf)
    pf.paragraph_format.space_after = Pt(6)


def add_abstract_block(doc, text: str):
    """초록 레이블 + 본문 — 모두 10pt Bold, 1단 전폭"""
    ph = doc.add_paragraph()
    ph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    rh = ph.add_run("초록")
    _apply_run_font(rh, size=PT10, bold=True)
    _set_single_spacing(ph)
    ph.paragraph_format.space_before = Pt(6)
    ph.paragraph_format.space_after  = Pt(2)

    pb = doc.add_paragraph()
    rb = pb.add_run(text)
    _apply_run_font(rb, size=PT10, bold=True)
    _set_single_spacing(pb)
    pb.paragraph_format.first_line_indent = Cm(0.5)
    pb.paragraph_format.space_after = Pt(4)


def add_heading(doc, text: str, level: int = 1):
    """본문 제목 11pt Bold — 위아래 한 줄 여백"""
    p = doc.add_paragraph()
    r = p.add_run(text)
    _apply_run_font(r, size=PT11, bold=True)
    _set_single_spacing(p)
    p.paragraph_format.space_before = Pt(10)   # 위 한 줄
    p.paragraph_format.space_after  = Pt(10)   # 아래 한 줄


def add_body(doc, text: str, indent: bool = False, space_after=Pt(3)):
    p = doc.add_paragraph()
    r = p.add_run(text)
    _apply_run_font(r, size=PT10)
    _set_single_spacing(p)
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.5)
    p.paragraph_format.space_after = space_after
    return p


def add_equation(
    doc,
    eq_text: str,
    number: str = None,
    space_after=Pt(4),
    align: str = "center",
    left_indent: float | None = None,
):
    """수식 — 중앙정렬, italic 수식 본문, 우측 수식 번호."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT if align == "left" else WD_ALIGN_PARAGRAPH.CENTER
    _set_single_spacing(p)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = space_after
    if left_indent is not None:
        p.paragraph_format.left_indent = Cm(left_indent)

    r_eq = p.add_run(eq_text)
    _apply_run_font(r_eq, size=PT10, italic=True)

    if number:
        r_num = p.add_run(f"        ({number})")
        _apply_run_font(r_num, size=PT10, italic=False)
    return p


def add_formula_block(doc, text: str, space_after=Pt(3), size=PT9):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.4)
    p.paragraph_format.right_indent = Cm(0.2)
    p.paragraph_format.space_after = space_after
    _set_single_spacing(p)
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = size
    rPr = r._r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), "Consolas")
    rFonts.set(qn("w:ascii"), "Consolas")
    rFonts.set(qn("w:hAnsi"), "Consolas")
    return p


def add_table(doc, tbl_data: dict):
    """표 제목 상단·왼쪽 정렬 / 위아래 한 줄 여백"""
    # 위 여백 (빈 줄 효과)
    gap_before = doc.add_paragraph()
    gap_before.paragraph_format.space_before = Pt(0)
    gap_before.paragraph_format.space_after  = Pt(0)
    gap_before.add_run().font.size = Pt(4)

    # 표 캡션 (왼쪽 정렬, 9pt Bold)
    cp = doc.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    cr = cp.add_run(tbl_data["caption"])
    _apply_run_font(cr, size=PT9, bold=True)
    _set_single_spacing(cp)
    cp.paragraph_format.space_before = Pt(0)
    cp.paragraph_format.space_after  = Pt(1)

    headers = tbl_data["headers"]
    rows    = tbl_data["rows"]
    n_cols  = len(headers)
    table_font = tbl_data.get("font_size", 9)
    tbl = doc.add_table(rows=1 + len(rows), cols=n_cols)
    tbl.style = "Table Grid"
    widths_cm = tbl_data.get("widths_cm")
    if widths_cm:
        tbl.autofit = False

    # 헤더 행
    hdr_cells = tbl.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        if widths_cm:
            hdr_cells[i].width = Cm(widths_cm[i])
        for p in hdr_cells[i].paragraphs:
            for r in p.runs:
                _apply_run_font(r, size=Pt(table_font), bold=True)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        tc = hdr_cells[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  "D9E1F2")
        tcPr.append(shd)

    # 데이터 행
    for ri, row_data in enumerate(rows):
        row_cells = tbl.rows[ri + 1].cells
        is_last = (ri == len(rows) - 1)
        for ci, cell_text in enumerate(row_data):
            row_cells[ci].text = cell_text
            if widths_cm:
                row_cells[ci].width = Cm(widths_cm[ci])
            for p in row_cells[ci].paragraphs:
                for r in p.runs:
                    _apply_run_font(r, size=Pt(table_font), bold=(is_last and ci == 0))
                p.alignment = (WD_ALIGN_PARAGRAPH.CENTER
                               if ci > 0 else WD_ALIGN_PARAGRAPH.LEFT)
        if is_last:
            for ci in range(n_cols):
                tc = tbl.rows[ri + 1].cells[ci]._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"),   "clear")
                shd.set(qn("w:color"), "auto")
                shd.set(qn("w:fill"),  "F2F2F2")
                tcPr.append(shd)

    # 아래 여백
    gap_after = doc.add_paragraph()
    gap_after.paragraph_format.space_before = Pt(0)
    gap_after.paragraph_format.space_after  = Pt(6)
    gap_after.add_run().font.size = Pt(4)


def add_figure(doc, img_path: str, caption: str, width_cm: float = 7.5):
    """그림 삽입 — 그림 하단 캡션, 왼쪽 정렬, 위아래 여백"""
    from docx.shared import Cm as _Cm

    gap_before = doc.add_paragraph()
    gap_before.paragraph_format.space_before = Pt(0)
    gap_before.paragraph_format.space_after  = Pt(0)
    gap_before.add_run().font.size = Pt(4)

    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_img.paragraph_format.space_before = Pt(0)
    p_img.paragraph_format.space_after  = Pt(1)
    run = p_img.add_run()
    run.add_picture(img_path, width=_Cm(width_cm))

    cp = doc.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    cr = cp.add_run(caption)
    _apply_run_font(cr, size=PT9, bold=True)
    _set_single_spacing(cp)
    cp.paragraph_format.space_before = Pt(0)
    cp.paragraph_format.space_after  = Pt(0)

    gap_after = doc.add_paragraph()
    gap_after.paragraph_format.space_before = Pt(0)
    gap_after.paragraph_format.space_after  = Pt(6)
    gap_after.add_run().font.size = Pt(4)


def add_equation_figure(doc, equations: list[str], caption: str, number: str = None):
    """수식 블록을 Figure처럼 배치하고 하단에 캡션을 둔다."""
    gap_before = doc.add_paragraph()
    gap_before.paragraph_format.space_before = Pt(0)
    gap_before.paragraph_format.space_after = Pt(0)
    gap_before.add_run().font.size = Pt(4)

    for i, eq in enumerate(equations):
        eq_indent = 0.6 if eq.rstrip().endswith(":") else 1.1
        add_equation(
            doc,
            eq,
            number=(number if i == 0 else None),
            space_after=Pt(2 if i < len(equations) - 1 else 1),
            align="left",
            left_indent=eq_indent,
        )

    cp = doc.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    cr = cp.add_run(caption)
    _apply_run_font(cr, size=PT9, bold=True)
    _set_single_spacing(cp)
    cp.paragraph_format.space_before = Pt(0)
    cp.paragraph_format.space_after = Pt(0)

    gap_after = doc.add_paragraph()
    gap_after.paragraph_format.space_before = Pt(0)
    gap_after.paragraph_format.space_after = Pt(6)
    gap_after.add_run().font.size = Pt(4)


def add_center_label(doc, text: str, space_after=Pt(1)):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_single_spacing(p)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = space_after
    r = p.add_run(text)
    _apply_run_font(r, size=PT10, bold=True)
    return p


def add_labeled_example_block(doc, title: str, lines, caption: str):
    """가운데 제목 + 라벨/값 계층형 예시 블록 + 하단 캡션."""
    gap_before = doc.add_paragraph()
    gap_before.paragraph_format.space_before = Pt(0)
    gap_before.paragraph_format.space_after = Pt(0)
    gap_before.add_run().font.size = Pt(4)

    add_center_label(doc, title, space_after=Pt(2))

    for label, value in lines:
        p_label = doc.add_paragraph()
        p_label.alignment = WD_ALIGN_PARAGRAPH.LEFT
        _set_single_spacing(p_label)
        p_label.paragraph_format.left_indent = Cm(0.6)
        p_label.paragraph_format.space_before = Pt(0)
        p_label.paragraph_format.space_after = Pt(0 if value is not None else 1)
        r_label = p_label.add_run(label)
        _apply_run_font(r_label, size=PT10, italic=True)

        if value is not None:
            if isinstance(value, str):
                value_lines = [value]
            else:
                value_lines = list(value)
            for idx, v in enumerate(value_lines):
                p_value = doc.add_paragraph()
                p_value.alignment = WD_ALIGN_PARAGRAPH.LEFT
                _set_single_spacing(p_value)
                p_value.paragraph_format.left_indent = Cm(1.1)
                p_value.paragraph_format.space_before = Pt(0)
                p_value.paragraph_format.space_after = Pt(1 if idx == len(value_lines) - 1 else 0)
                r_value = p_value.add_run(v)
                _apply_run_font(r_value, size=PT10, italic=True)

    cp = doc.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    cr = cp.add_run(caption)
    _apply_run_font(cr, size=PT9, bold=True)
    _set_single_spacing(cp)
    cp.paragraph_format.space_before = Pt(0)
    cp.paragraph_format.space_after = Pt(0)

    gap_after = doc.add_paragraph()
    gap_after.paragraph_format.space_before = Pt(0)
    gap_after.paragraph_format.space_after = Pt(6)
    gap_after.add_run().font.size = Pt(4)


def add_labeled_example_table(doc, title: str, lines, caption: str):
    """1x1 표 안에 제목/라벨/값 예시를 넣고 하단 캡션을 붙인다."""
    gap_before = doc.add_paragraph()
    gap_before.paragraph_format.space_before = Pt(0)
    gap_before.paragraph_format.space_after = Pt(0)
    gap_before.add_run().font.size = Pt(4)

    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    table.autofit = True
    cell = table.cell(0, 0)

    # top padding paragraph
    p_pad = cell.paragraphs[0]
    p_pad.paragraph_format.space_before = Pt(0)
    p_pad.paragraph_format.space_after = Pt(0)
    p_pad.add_run().font.size = Pt(2)

    p_title = cell.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_single_spacing(p_title)
    p_title.paragraph_format.space_before = Pt(0)
    p_title.paragraph_format.space_after = Pt(2)
    r_title = p_title.add_run(title)
    _apply_run_font(r_title, size=PT10, bold=True)

    for label, value in lines:
        p_label = cell.add_paragraph()
        p_label.alignment = WD_ALIGN_PARAGRAPH.LEFT
        _set_single_spacing(p_label)
        p_label.paragraph_format.left_indent = Cm(0.35)
        p_label.paragraph_format.space_before = Pt(0)
        p_label.paragraph_format.space_after = Pt(0 if value is not None else 1)
        r_label = p_label.add_run(label)
        _apply_run_font(r_label, size=PT10, italic=True)

        if value is not None:
            value_lines = [value] if isinstance(value, str) else list(value)
            for idx, v in enumerate(value_lines):
                p_value = cell.add_paragraph()
                p_value.alignment = WD_ALIGN_PARAGRAPH.LEFT
                _set_single_spacing(p_value)
                p_value.paragraph_format.left_indent = Cm(0.85)
                p_value.paragraph_format.space_before = Pt(0)
                p_value.paragraph_format.space_after = Pt(1 if idx == len(value_lines) - 1 else 0)
                r_value = p_value.add_run(v)
                _apply_run_font(r_value, size=PT10, italic=True)

    cp = doc.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    cr = cp.add_run(caption)
    _apply_run_font(cr, size=PT9, bold=True)
    _set_single_spacing(cp)
    cp.paragraph_format.space_before = Pt(1)
    cp.paragraph_format.space_after = Pt(0)

    gap_after = doc.add_paragraph()
    gap_after.paragraph_format.space_before = Pt(0)
    gap_after.paragraph_format.space_after = Pt(6)
    gap_after.add_run().font.size = Pt(4)


def make_yolo_roi_figure(save_path: str):
    """
    3-panel single-row figure:
    [Class A reference] [Class B reference] [Test sample: GT=A, CNN→B✗, YOLO→A✓]
    """
    import numpy as np
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    from matplotlib.patches import Rectangle, Circle

    matplotlib.rcParams.update({"font.family": "DejaVu Sans", "font.size": 8})

    def _chip_pixel_points(kind):
        if kind == "star":
            return {
                (2, 0),
                (1, 1), (2, 1), (3, 1),
                (0, 2), (1, 2), (2, 2), (3, 2), (4, 2),
                (1, 3), (2, 3), (3, 3),
                (2, 4),
            }
        if kind == "x":
            return {
                (0, 0), (4, 0),
                (1, 1), (3, 1),
                (2, 2),
                (1, 3), (3, 3),
                (0, 4), (4, 4),
            }
        return set()

    def _wafer_defect_chip_points():
        return {
            (-2, -2), (-1, -2), (0, -2), (1, -2),
            (-2, -1), (-1, -1), (0, -1), (1, -1), (2, -1),
            (-2, 0), (-1, 0), (0, 0), (1, 0), (2, 0),
            (-2, 1), (-1, 1), (0, 1), (1, 1),
        }

    def _draw_defect_chip(ax, x0, y0, size, kind):
        ax.add_patch(Rectangle((x0, y0), size, size, facecolor="#F4F4F4", edgecolor="#666666", lw=0.35))
        N = 5
        cell = size / N
        pts = _chip_pixel_points(kind)
        for y in range(N):
            for x in range(N):
                color = "#D94040" if (x, y) in pts else "#C8E0C8"
                ax.add_patch(Rectangle(
                    (x0 + x * cell + 0.12 * cell, y0 + (N - 1 - y) * cell + 0.12 * cell),
                    cell - 0.24 * cell, cell - 0.24 * cell,
                    facecolor=color, edgecolor="#DDDDDD", lw=0.12))
        ax.add_patch(Rectangle((x0, y0), size, size, fill=False, ec="#444444", lw=0.8))

    def _draw_normal_chip(ax, x0, y0, size):
        ax.add_patch(Rectangle(
            (x0 + 0.08 * size, y0 + 0.08 * size), size - 0.16 * size, size - 0.16 * size,
            facecolor="#C8E0C8", edgecolor="#AAAAAA", lw=0.18))

    def _draw_wafer(ax, kind, highlight=False):
        S = 100
        chip = 6
        cx, cy, R = 33, 56, 27
        center_pts = _wafer_defect_chip_points()
        for xi in range(0, S, chip):
            for yi in range(0, S, chip):
                ccx, ccy = xi + chip // 2, yi + chip // 2
                d = np.sqrt((ccx - cx) ** 2 + (ccy - cy) ** 2)
                if d > R:
                    continue
                gx = round((ccx - cx) / chip)
                gy = round((ccy - cy) / chip)
                if (gx, gy) in center_pts:
                    x0, y0, size = xi + 0.5, yi + 0.5, chip - 1.3
                    _draw_defect_chip(ax, x0, y0, size, kind)
                    if highlight:
                        ax.add_patch(Rectangle(
                            (x0 - 0.35, y0 - 0.35), size + 0.7, size + 0.7,
                            fill=False, ec="#00A82D", lw=0.9))
                else:
                    _draw_normal_chip(ax, xi + 0.5, yi + 0.5, chip - 1.3)
        ax.add_patch(Circle((cx, cy), R, fill=False, ec="#333333", lw=1.4))
        ax.add_patch(Rectangle((22, 45), 22, 22, fill=False, ec="#FFB300", lw=1.8, ls="--"))

    def _draw_panel(ax, pattern, title, pred_lines=None, highlight=False):
        ax.set_xlim(0, 100)
        ax.set_ylim(0, 100)
        ax.set_aspect("equal")
        ax.axis("off")
        ax.set_title(title, fontsize=9, fontweight="bold", pad=4)
        _draw_wafer(ax, pattern, highlight=highlight)
        ax.annotate("", xy=(41, 56), xytext=(70, 58),
                    arrowprops=dict(arrowstyle="->", lw=1.0, color="#777777"))
        _draw_defect_chip(ax, x0=66, y0=46, size=24, kind=pattern)
        ax.text(78, 43, "chip pixel pattern", ha="center", va="top",
                fontsize=7.3, color="#555555", style="italic")
        if pred_lines:
            for y, text, color in pred_lines:
                ax.text(50, y, text, ha="center", va="center",
                        fontsize=8, fontweight="bold", color=color)

    fig, axes = plt.subplots(1, 3, figsize=(8.6, 3.0),
                             facecolor="white",
                             gridspec_kw={"wspace": 0.10})

    _draw_panel(axes[0], "star", "Class A  (object star mark)")
    _draw_panel(axes[1], "x", "Class B  (object x mark)")
    _draw_panel(
        axes[2], "star", "Test sample  (GT: Class A)",
        pred_lines=[
            (16, "CNN pred: Class B  \u2717  (conf 0.54)", "#CC0000"),
            (7,  "YOLO ROI pred: Class A  \u2713", "#007700"),
        ],
        highlight=True,
    )

    axes[2].add_patch(Rectangle((20, 66), 28, 8, facecolor="#00A82D", edgecolor="none"))
    axes[2].text(34, 70, "star mark 0.91", ha="center", va="center",
                 fontsize=8, fontweight="bold", color="white")

    # 범례 (우상단 inset)
    ax3 = axes[2]
    for dy, fc, label in [(0.92, "#C8E0C8", "Normal"),
                           (0.78, "#D94040", "Defect")]:
        ax3.add_patch(Rectangle((0.68, dy - 0.03), 0.08, 0.09,
                                  facecolor=fc, edgecolor="#888888", lw=0.5,
                                  transform=ax3.transAxes))
        label_text = "Normal chip" if label == "Normal" else "Defect pixel"
        ax3.text(0.79, dy + 0.015, label_text, transform=ax3.transAxes,
                 fontsize=6.5, va="center", color="#333333")

    fig.text(0.08, 0.03, "wafer-level: same center distribution (CNN confused)",
             fontsize=10, color="#777777", style="italic")

    fig.savefig(save_path, dpi=150, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    return save_path


def make_data_conversion_figure(save_path: str):
    """
    2-panel figure:
    (1) hex -> decimal -> grade conversion
    (2) RGB repetition -> palette index compression
    """
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    from matplotlib.patches import FancyBboxPatch, FancyArrowPatch

    matplotlib.rcParams.update({"font.family": "DejaVu Sans", "font.size": 8})

    fig, axes = plt.subplots(1, 2, figsize=(8.4, 2.8), facecolor="white")
    for ax in axes:
        ax.set_xlim(0, 10)
        ax.set_ylim(0, 6)
        ax.axis("off")

    # Panel 1: hex conversion
    ax = axes[0]
    ax.text(5, 5.55, "Hex to Grade conversion", ha="center", va="center",
            fontsize=9, fontweight="bold", color="#111827")
    panel = FancyBboxPatch((0.4, 0.7), 9.2, 4.2,
                           boxstyle="round,pad=0.02,rounding_size=0.08",
                           linewidth=1.0, edgecolor="#CBD5E1",
                           facecolor="#F8FAFC")
    ax.add_patch(panel)
    ax.text(5, 4.25, "0A  09  0B  0C", ha="center", va="center",
            fontsize=10, fontweight="bold", color="#1D4ED8")
    ax.add_patch(FancyArrowPatch((5, 3.95), (5, 3.25), arrowstyle="-|>",
                                 mutation_scale=12, lw=1.2, color="#64748B"))
    ax.text(5, 2.95, "10   9  11  12", ha="center", va="center",
            fontsize=10, fontweight="bold", color="#7C3AED")
    ax.add_patch(FancyArrowPatch((5, 2.65), (5, 1.95), arrowstyle="-|>",
                                 mutation_scale=12, lw=1.2, color="#64748B"))
    ax.text(5, 1.65, "1   0   2   3", ha="center", va="center",
            fontsize=10, fontweight="bold", color="#059669")
    ax.text(5, 0.95, "v = int(hex,16),  g = 0 (v=0),  g = v - 9 (v>0)",
            ha="center", va="center", fontsize=7.2, color="#475569")

    # Panel 2: palette indexed png
    ax = axes[1]
    ax.text(5, 5.55, "Palette-indexed PNG", ha="center", va="center",
            fontsize=9, fontweight="bold", color="#111827")
    panel = FancyBboxPatch((0.4, 0.7), 9.2, 4.2,
                           boxstyle="round,pad=0.02,rounding_size=0.08",
                           linewidth=1.0, edgecolor="#CBD5E1",
                           facecolor="#F8FAFC")
    ax.add_patch(panel)
    ax.text(5, 4.35,
            "(123,54,24)  (123,54,24)  (123,54,24)  (123,54,24)",
            ha="center", va="center", fontsize=7.6,
            fontweight="bold", color="#92400E")
    ax.add_patch(FancyArrowPatch((5, 4.0), (5, 3.2), arrowstyle="-|>",
                                 mutation_scale=12, lw=1.2, color="#64748B"))
    ax.text(5, 2.95, "palette[1] = (123,54,24)", ha="center", va="center",
            fontsize=8.4, fontweight="bold", color="#7C3AED")
    ax.add_patch(FancyArrowPatch((5, 2.65), (5, 1.85), arrowstyle="-|>",
                                 mutation_scale=12, lw=1.2, color="#64748B"))
    ax.text(5, 1.55, "1   1   1   1", ha="center", va="center",
            fontsize=11, fontweight="bold", color="#059669")
    ax.text(5, 0.95,
            "RGB values are defined once in the palette; pixels store only the index.",
            ha="center", va="center", fontsize=7.0, color="#475569")

    fig.savefig(save_path, dpi=160, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    return save_path


def make_ui_figure(save_path: str):
    """
    3-panel figure: Failbit Map | FTN overlay | QTN overlay
    Failbit Map = spatial defect result
    FTN/QTN = chip-level electrical Measure for root-cause analysis
    """
    import numpy as np
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    from matplotlib.patches import Rectangle, Circle
    from matplotlib.colors import Normalize
    import matplotlib.cm as cm

    matplotlib.rcParams.update({"font.family": "DejaVu Sans", "font.size": 8})

    np.random.seed(7)
    S, chip = 100, 6
    cx, cy, R = S // 2, S // 2, S // 2 - 4

    chips = []
    for xi in range(0, S, chip):
        for yi in range(0, S, chip):
            ccx, ccy = xi + chip // 2, yi + chip // 2
            d = np.sqrt((ccx - cx) ** 2 + (ccy - cy) ** 2)
            if d > R:
                continue
            is_defect = (d < R * 0.36) and (np.random.random() > 0.12)
            ftn = int(np.random.normal(220, 30)) if is_defect else int(np.random.normal(12, 5))
            ftn = max(0, ftn)
            qtn = int(np.random.normal(160, 25)) if is_defect else int(np.random.normal(20, 8))
            qtn = max(0, qtn)
            chips.append((xi, yi, is_defect, ftn, qtn))

    ftn_max = max(c[3] for c in chips)
    qtn_max = max(c[4] for c in chips)
    ftn_norm = Normalize(vmin=0, vmax=ftn_max)
    qtn_norm = Normalize(vmin=0, vmax=qtn_max)

    fig, axes = plt.subplots(1, 3, figsize=(8.2, 2.5),
                             facecolor="white",
                             gridspec_kw={"wspace": 0.18})

    titles    = ["Failbit Map", "FTN overlay", "QTN overlay"]
    subtitles = ["Spatial defect pattern", "Fail Total Number / chip", "Qualified Test Number / chip"]

    for ax_idx, ax in enumerate(axes):
        ax.add_patch(Circle((cx, cy), R, fill=False, ec="#333333", lw=1.5))
        for (xi, yi, is_defect, ftn, qtn) in chips:
            if ax_idx == 0:
                color = "#D94040" if is_defect else "#C8E0C8"
            elif ax_idx == 1:
                color = cm.Reds(ftn_norm(ftn))
            else:
                color = cm.Oranges(qtn_norm(qtn))
            ax.add_patch(Rectangle(
                (xi + 0.5, yi + 0.5), chip - 1.4, chip - 1.4,
                facecolor=color, edgecolor="#AAAAAA", lw=0.2))
        ax.set_xlim(0, S); ax.set_ylim(0, S)
        ax.set_aspect("equal"); ax.axis("off")
        ax.set_title(titles[ax_idx], fontsize=9, fontweight="bold", pad=4)
        ax.text(0.5, -0.04, subtitles[ax_idx], transform=ax.transAxes,
                ha="center", fontsize=7.5, color="#555555", style="italic")

    fig.savefig(save_path, dpi=150, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    return save_path


def make_cluster_figure(save_path: str):
    """
    Unknown defect clustering 결과 모식도
    3 clusters × 6 mini-wafer thumbnails
    유사 패턴끼리 자동 그룹핑됨을 시각적으로 보여줌
    """
    import numpy as np
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    from matplotlib.patches import Rectangle, Circle

    matplotlib.rcParams.update({"font.family": "DejaVu Sans", "font.size": 7})

    def _draw_mini(ax, pattern, seed, bg):
        np.random.seed(seed)
        S, chip = 50, 5
        cx, cy, R = S // 2, S // 2, S // 2 - 2
        ax.set_facecolor(bg)
        for xi in range(0, S, chip):
            for yi in range(0, S, chip):
                ccx, ccy = xi + chip // 2, yi + chip // 2
                d = np.sqrt((ccx - cx) ** 2 + (ccy - cy) ** 2)
                if d > R:
                    continue
                color = "#C8E0C8"
                if pattern == "partial_ring":
                    if 0.40 * R < d < 0.68 * R:
                        ang = np.arctan2(ccy - cy, ccx - cx)
                        if not (-0.4 < ang < 1.3):
                            color = "#D94040" if np.random.random() > 0.10 else color
                elif pattern == "cross":
                    if (abs(ccx - cx) < chip * 0.85 or abs(ccy - cy) < chip * 0.85) and d < R * 0.88:
                        color = "#D94040" if np.random.random() > 0.10 else color
                elif pattern == "diagonal":
                    dist = abs((ccx - cx) - (ccy - cy)) / 1.414
                    if dist < chip * 0.82 and d < R * 0.90:
                        color = "#D94040" if np.random.random() > 0.10 else color
                ax.add_patch(Rectangle(
                    (xi + 0.3, yi + 0.3), chip - 0.8, chip - 0.8,
                    facecolor=color, edgecolor="#CCCCCC", lw=0.15))
        ax.add_patch(Circle((cx, cy), R, fill=False, ec="#555555", lw=0.7))
        ax.set_xlim(0, S); ax.set_ylim(0, S)
        ax.set_aspect("equal"); ax.axis("off")

    clusters = [
        ("Cluster 1  (12 wafers)", "partial_ring", [11, 22, 35, 40, 53, 64]),
        ("Cluster 2  (9 wafers)",  "cross",        [14, 27, 31, 46, 57, 70]),
        ("Cluster 3  (8 wafers)",  "diagonal",     [13, 24, 38, 43, 56, 67]),
    ]
    bg_colors = ["#FFF5F5", "#F5F0FF", "#F5FFF8"]

    fig, axes = plt.subplots(2, 9, figsize=(8.2, 2.2),
                             facecolor="white",
                             gridspec_kw={"wspace": 0.07, "hspace": 0.07})
    plt.subplots_adjust(left=0.01, right=0.99, top=0.80, bottom=0.02)

    col_offsets = [0, 3, 6]
    for ci, (label, pattern, seeds) in enumerate(clusters):
        c0 = col_offsets[ci]
        for i, seed in enumerate(seeds):
            row, col = divmod(i, 3)
            _draw_mini(axes[row][c0 + col], pattern, seed, bg_colors[ci])
        # 클러스터 레이블 — 가운데 열 위
        axes[0][c0 + 1].set_title(label, fontsize=7.5, fontweight="bold",
                                   pad=3, color="#222222")

    fig.savefig(save_path, dpi=150, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    return save_path


def add_subheading(doc, text: str):
    """소절 제목 — 11pt Bold, 위 5pt / 아래 3pt (2-page 압축용)"""
    p = doc.add_paragraph()
    r = p.add_run(text)
    _apply_run_font(r, size=PT11, bold=True)
    _set_single_spacing(p)
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after  = Pt(3)


def add_bullet(doc, text: str):
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(text)
    _apply_run_font(r, size=PT10)
    _set_single_spacing(p)
    p.paragraph_format.space_after = Pt(1)


def add_refs(doc, refs: list):
    add_heading(doc, "참고문헌", level=1)
    for ref in refs:
        p = doc.add_paragraph()
        r = p.add_run(ref)
        _apply_run_font(r, size=PT9)
        _set_single_spacing(p)
        p.paragraph_format.space_after      = Pt(1)
        p.paragraph_format.left_indent      = Cm(0.5)
        p.paragraph_format.first_line_indent = Cm(-0.5)


# ══════════════════════════════════════════════════════════════
# 상세 버전 (paper_detailed.docx)
# ══════════════════════════════════════════════════════════════
def build_detailed() -> Document:
    doc = Document()

    # ── 최종 섹션(본문) 설정: A4, 여백, 2단 ──────────────────
    sec = doc.sections[0]
    sec.page_width  = Cm(21)
    sec.page_height = Cm(29.7)
    set_margins(sec)
    set_two_columns(sec, space_cm=0.5)
    doc.styles["Normal"].font.size = PT10
    doc.styles["Normal"].font.name = F_KO

    # ── 1단 영역: 제목 / 저자 / 초록 ─────────────────────────
    add_title_block(doc, TITLE_KO_CLAUDE, TITLE_EN_CLAUDE)
    add_author_block(doc, AUTHORS, AFFIL)
    add_abstract_block(doc, ABSTRACT_CLAUDE)

    # 1단→2단 전환 (연속 섹션 구분)
    add_single_to_double_break(doc)

    # ── 2단 영역: 본문 ────────────────────────────────────────

    # 1. 서론
    add_heading(doc, "1. 서론", level=1)
    add_body(doc,
        "웨이퍼 Failbit Map은 메모리 테스트 결과를 공간적으로 표현한 데이터로서, 불량의 위치, "
        "분포, 방향성, 밀집도와 같은 정보를 직접적으로 담고 있다. center, edge ring, local, "
        "scratch 계열 패턴은 단순 시각적 모양이 아니라 공정 이상 원인과 연결되는 공간적 의미를 "
        "가지므로, 이를 자동 분류하고 해석하는 기술은 반도체 제조 현장에서 높은 실용성을 갖는다 [1][2]. "
        "그러나 현업에서 wafer map 분석은 더욱 근본적인 어려움을 안고 있다. 수치 계측값(yield, "
        "BIN 분포)은 집계가 가능하지만, 불량의 공간적 패턴은 이미지를 직접 보아야만 파악된다. "
        "하루에 발생하는 wafer 수 대비 맵을 대량으로 열람할 수 있는 인프라가 없어, 실무자들은 "
        "관심 있는 wafer의 맵 파일을 개별적으로 저장하고 메신저나 이메일로 공유하는 방식에 "
        "의존해 왔다. 이는 이상 징후를 놓치기 쉽고, 분석 결과가 개인에게 귀속되어 조직 차원의 "
        "지식으로 축적되지 않는 구조적 문제를 낳는다.", indent=True)

    add_body(doc,
        "본 시스템 설계 과정에서 실무 요구사항 조사를 통해 세 가지 핵심 과제가 도출되었다. "
        "첫째, 불량 라벨 등록은 해당 공정에 정통한 특정 전문가만 수행할 수 있어 데이터 확보 자체가 "
        "어렵고, 기존 방식은 라벨링 절차가 불편하여 전문가의 부담이 크다. 이에 라벨링 부담을 줄이고 "
        "비전문가도 검토에 참여할 수 있는 UI 기반 라벨 확장 워크플로우가 요구되었다. "
        "둘째, 공정 변화나 신규 제품 도입 시 기존 클래스 체계에 없는 unknown defect가 갑작스럽게 "
        "발생하는 경우가 있으며, 조기에 포착하지 못하면 수율 손실로 이어질 수 있다 [4]. "
        "셋째, wafer map과 chip 계측값을 함께 조회하며 종합적으로 분석할 수 있는 통합 UI에 대한 "
        "요구가 있었다.", indent=True)

    add_body(doc,
        "이에 본 연구는 단일 분류 모델의 성능 개선을 넘어서 데이터 생성부터 등록 불량 분류, "
        "미등록 불량 군집화, UI 기반 검토와 라벨 확장까지 연결되는 통합 구조를 설계하였다. "
        "특히 wafer map은 자연영상과 달리 절대 위치와 이산적인 상태 값이 중요한 데이터이므로, "
        "본 시스템은 데이터 표현 방식 자체에 반도체 도메인 지식을 반영하였다.", indent=True)

    # 2. 제안 방법
    add_heading(doc, "2. 제안 방법", level=1)

    add_heading(doc, "2.1 데이터 생성 및 표현 계층", level=2)
    add_body(doc,
        "본 계층의 목적은 단순히 이미지를 생성하는 것이 아니라, wafer의 공간 패턴 정보와 chip-level "
        "계측 정보를 동일 좌표계에서 결합한 공통 표현을 구축하는 데 있다. Bucket A의 raw fail-map "
        "로그는 불량의 공간 분포를 정밀하게 담고 있지만 chip-level measurement를 포함하지 않고, "
        "Bucket B의 계측 로그는 BIN, FBT, QVL과 같은 수치 정보를 제공하지만 공간 패턴 자체를 직접 "
        "보여주지 않는다. 두 소스를 정합하여 wafer image와 chip annotation 파일을 동시에 생성하면, 동일 "
        "wafer에서 chip label 생성, YOLO object detection용 ROI/annotation 관리, 사용자 분석 UI "
        "제공, 향후 이미지-계측 multimodal 학습에 이르기까지 하나의 공통 데이터 자산으로 재사용할 수 있다.")
    add_body(doc,
        "구현 측면에서 중요한 제약은 두 로그가 동일 wafer를 가리키더라도 파일명 형식이 일치하지 않고 "
        "저장 시각에도 편차가 존재한다는 점이다. 이 때문에 단순 파일명 join이나 사후 DB 매칭에 "
        "의존하면 전수 검색 비용이 커지고, 실시간에 가까운 데이터 생성 파이프라인을 구성하기 어렵다. "
        "따라서 본 시스템은 wafer 식별 규칙과 ±10초 시간 오프셋을 이용해 Bucket A의 primary fail map "
        "raw 파일과 Bucket B의 chip-level measurement 로그를 로그 단계에서 직접 정합한다. 이후 "
        "Bucket A 로그의 `X=, Y=, b=` 헤더와 뒤따르는 hex grade block을 Cython 기반 Hex 파싱으로 "
        "변환하여 chip tile array를 복원하고, Bucket B에서 읽은 BIN, FBT, QVL 및 wafer-level yield, "
        "sys, lt, tm을 chip annotation 파일에 병합한다. 이 결과 분류기와 UI는 동일한 wafer 좌표계와 동일한 "
        "chip 인덱스를 공유하게 된다.")
    add_body(doc,
        "wafer map 저장 포맷으로 32-color 8-bit palette-indexed PNG를 채택한 이유도 도메인 지식에 "
        "기반한다. 실제 생성 이미지는 자연영상처럼 수천~수만 색을 쓰지 않고, grade, background, "
        "text, border, BIN 등 검사 결과를 표현하는 약 20개 내외의 이산 색상으로 구성된다. "
        "따라서 JPEG 같은 손실 압축은 불필요하고, RGB PNG처럼 픽셀에 색값을 직접 저장할 필요도 없다. "
        "palette PNG는 픽셀에 '의미 인덱스'를 저장하고 색상은 PLTE에서만 정의하므로, IDAT 구조를 "
        "유지한 채 PLTE만 교체하여 사용자별 개인 색상 scheme을 즉시 적용할 수 있다. "
        "24bit RGB 대비 약 65% 파일 크기 절감과 화질 완전 보존을 달성한다.")

    add_heading(doc, "2.2 등록 불량 분류 경로", level=2)
    add_body(doc,
        "등록 불량 분류는 16개 클래스 closed-set 문제로 설정하였다. 웨이퍼 맵은 자연영상처럼 "
        "경계(edge)가 명확하지 않고, 전기적 노이즈가 공간적으로 뭉쳐진 형태이다. "
        "Fine-Grained Visual Classification(FGVC) 계열 접근법을 시도하였으나, 자연영상의 텍스처·경계 "
        "구조를 가정한 방법들은 성능 향상으로 이어지지 않았다. Focal Loss 및 class weight 조정으로 "
        "클래스 불균형 완화를 시도하였으나 효과는 제한적이었다.")
    add_body(doc,
        "성능 향상의 실질적 경로는 두 단계로 정리된다. 첫 단계는 백본 선택과 하이퍼파라미터 최적화이다. "
        "EfficientNetV2, MaxViT, ViT, Swin Transformer 등을 Hugging Face timm 라이브러리를 통해 "
        "동일 조건에서 fine-tuning 비교하였다. 총 약 1500개(validation 약 500개) 규모의 데이터에서 "
        "self-attention 학습에 데이터가 부족한 ViT 계열보다, ConvNeXt V2에서 제안된 FCMAE "
        "(Fully Convolutional Masked Autoencoder) 기반 사전학습과 384×384 입력 해상도를 지원하는 "
        "ConvNeXtV2-Base가 더 적합하였다. 공식 결과에서도 ConvNeXt V2-Base 384 모델은 "
        "ImageNet-1K 기준 87.7% top-1 accuracy를 보고하며 [5], 본 연구의 backbone 비교에서도 "
        "val F1 0.92로 1위를 기록하였다. "
        "이후 Optuna 기반 하이퍼파라미터 탐색(LR, weight decay, scheduler, dropout, label smoothing, "
        "augmentation, batch size)까지 포함한 결과이다.")
    add_table(doc, TABLE_BACKBONE)
    add_body(doc,
        "두 번째 단계는 2-stage ROI + YOLO 구조이다. 모든 샘플에 YOLO를 적용하는 방식은 채택하지 "
        "않았다. 이유는 두 가지이다. 첫째, 연산 비용—매 샘플마다 object detection을 수행하면 "
        "처리량이 급격히 감소한다. 둘째, wafer 불량 패턴은 전역(global) 맥락 기반이다. center, "
        "edge_ring, near_full 같은 패턴은 웨이퍼 전체 맵에서의 분포를 봐야 구분되며, chip 단위 "
        "local detection만으로는 판단이 불가능하다. 따라서 CNN은 global 패턴 판단, YOLO는 "
        "borderline 케이스의 local 검증을 담당하는 역할 분리 구조를 설계하였다.")
    add_body(doc,
        "ROI 보강은 classifier confidence < 0.80 또는 precision/recall < 0.80인 difficult class일 때만 "
        "선택적으로 수행한다. 특히 낮은 confidence 구간은 실제 오분류 비율이 높은 영역이므로, "
        "전역 패턴만으로는 class 분리가 충분하지 않은 신호로 해석하였다. Grad-CAM [6] dynamic ROI는 "
        "클래스별 spatial prior와 α-블렌딩하여 heatmap 불안정성을 보완한다. YOLO 학습에서는 손실 "
        "항목 중 cls_loss의 영향이 가장 컸다. "
        "ROI 내부에는 defect object가 1~2개 수준으로 존재하므로 box와 objectness 항은 빠르게 수렴하고, "
        "scratch·local·arc·edge 계열처럼 형태가 유사한 객체를 어떤 class로 구분하느냐가 핵심이다. "
        "cls gain을 강화하였을 때 ROI correction precision과 최종 F1이 안정적으로 상승하였다.")
    add_table(doc, TABLE_THRESH)
    add_body(doc,
        "CNN과 YOLO 모두 좌우·상하 반전 증강을 비활성화하였다. 반전은 edge_loc 발생 방향, scratch "
        "진행 방향 등 공간 방향 정보를 파괴한다. 동일한 외형이라도 left/right edge_loc는 설비 "
        "레이아웃 기준으로 서로 다른 공정 원인을 가지므로, flip을 허용하면 이 구분이 불가능해진다.")
    add_table(doc, TABLE1)
    add_table(doc, TABLE2A)
    add_table(doc, TABLE2)

    add_heading(doc, "2.3 미등록 불량 경로", level=2)
    add_body(doc,
        "unknown defect에 대해서는 ConvNeXtV2 backbone 기반 자기지도 대조 학습과 HDBSCAN [8] "
        "군집화를 결합하였다.")
    add_body(doc,
        "[손실 함수 선택] 초기에 Supervised Contrastive Learning(SupCon) [3]을 시험하였다. "
        "SupCon은 기등록 불량 군집 밀집도는 향상시켰으나, 라벨 없는 unknown defect 영역의 임베딩 "
        "분리도가 오히려 저하되었다. 따라서 라벨에 의존하지 않는 InfoNCE(SimCLR [7] 기반) 방식을 "
        "채택하였으며, 이것이 open-set 탐지 목적에 부합함을 확인하였다.")
    add_body(doc,
        "[Train/Test 분리 오해 해소] contrastive learning은 정답을 예측하는 것이 아니라 유사한 "
        "샘플을 임베딩 공간에서 인접하게 배치하는 학습이다. 과적합 리스크가 없으므로 전체 데이터를 "
        "학습에 활용할 수 있으며, 이 방식이 임베딩 공간의 커버리지를 높여 군집 품질 향상으로 이어졌다.")
    add_body(doc,
        "[Local InfoNCE 샘플링] 피처 맵을 6×6 격자(grid36_full)로 분할하고, 중심·외곽 주요 셀에 "
        "더 많은 앵커를 배치하는 structured sampling을 적용하였다. wafer map에서 위치 자체가 공정 "
        "의미를 가지기 때문이다.")
    add_body(doc,
        "[flip 미사용] 동일한 외형이라도 좌우 반전된 불량은 공정 원인이 다를 수 있다. "
        "Global InfoNCE(Queue 16K, false-negative 마스킹 sim≥0.72) + Local InfoNCE + "
        "HDBSCAN quality filter(size≥12, prob≥0.55, persist≥0.20) 조합으로 cluster 안정성을 확보하고, "
        "medoid representative를 저장하여 전문가 검토 부담을 줄였다.")
    add_table(doc, TABLE_HDBSCAN)

    add_heading(doc, "2.4 분석 UI와 multimodal 준비", level=2)
    add_body(doc,
        "mapviewer는 wafer map 의미 인덱스 구조와 chip-level 계측값을 함께 해석하는 분석 "
        "인터페이스이다. 개인 색상 scheme, chip annotation, composite map, BIN/FBT/QVL overlay를 "
        "제공하며, chip annotation 파일을 기반으로 chip geometry와 measurement를 동일 좌표계에 정합한다. "
        "BIN, FBT, QVL은 단순 UI 표시용 값이 아니라, 이미지와 정합 저장되어 향후 "
        "multimodal 학습 입력으로 직접 활용할 수 있는 데이터 자산으로 확보하였다.")

    # 3. 실험 결과
    add_heading(doc, "3. 실험 결과 및 논의", level=1)
    add_body(doc,
        "등록 불량 경로에서는 Optuna 기반 1-stage optimization으로 weighted F1 0.92를 달성하였고, "
        "ROI enhancement와 YOLO 보정을 통해 0.95까지 향상되었다. edge_loc, edge_ring, local 계열처럼 "
        "공간적으로 유사한 패턴 간 혼동과, center 부근 동일 영역 패턴 간 혼동이 줄어드는 효과가 "
        "확인되었다. difficult class 및 low-confidence 케이스에 한해 선택 적용한 YOLO 보정의 "
        "correction precision은 90% 이상을 기록하였다.")
    add_body(doc,
        "미등록 불량은 사전 Ground Truth 구성이 어려워 세 층위의 평가를 적용한다.")
    add_bullet(doc, "(1) 내부 지표: Silhouette Score 및 HDBSCAN 연성 멤버십 확률 분포")
    add_bullet(doc, "(2) 전문가 동의율: 각 군집을 담당 엔지니어가 유효한 신규 불량으로 확정한 비율")
    add_bullet(doc, "(3) 하위 태스크 전이 성능: 군집 라벨 등록 후 재학습 시 F1 향상 폭")
    add_table(doc, TABLE_COMPARE)

    # 4. 결론
    add_heading(doc, "4. 결론", level=1)
    add_body(doc,
        "본 연구는 반도체 웨이퍼 불량 분석을 위해 데이터 생성, 등록 불량 분류, 미등록 불량 군집화, "
        "분석 UI를 하나의 파이프라인으로 통합하였다. 약 20개 내외 이산 색상 특성을 활용한 "
        "palette PNG 표현, ConvNeXtV2 + Optuna로 F1 0.92, 선택적 ROI + YOLO 보정으로 0.95, "
        "InfoNCE 기반 대조 학습 + HDBSCAN으로 unknown defect 자동 군집화, BIN/FBT/QVL 정합 저장으로 "
        "multimodal 확장 기반을 마련하였다. 향후 칩 레벨 다중 라벨 분류, 이미지+계측 multimodal "
        "모델, 전체 제품군 확대를 진행할 예정이다.", indent=True)

    add_refs(doc, REFS)
    return doc


# ══════════════════════════════════════════════════════════════
# 2페이지 압축 버전 (paper_2page.docx)
# ══════════════════════════════════════════════════════════════
def build_2page() -> Document:
    doc = Document()

    sec = doc.sections[0]
    sec.page_width  = Cm(21)
    sec.page_height = Cm(29.7)
    set_margins(sec)
    set_two_columns(sec, space_cm=0.5)
    doc.styles["Normal"].font.size = PT10
    doc.styles["Normal"].font.name = F_KO

    # ── 1단: 제목 / 저자 / 초록 ──────────────────────────────
    add_title_block(doc, TITLE_KO_CODEX, TITLE_EN_CODEX)
    add_author_block(doc, AUTHORS, AFFIL)
    add_abstract_block(doc, ABSTRACT_CODEX)
    add_single_to_double_break(doc)

    # ── 2단: 본문 ─────────────────────────────────────────────
    add_heading(doc, "1. 서론", level=1)
    add_body(doc,
        "반도체 Failbit Map 불량 패턴 분류는 수율 관리의 핵심 과제이나, 맵 이미지를 대량으로 열람할 "
        "수 있는 인프라가 없어 실무자들이 개별 저장·공유 방식에 의존하는 구조적 문제가 있다. "
        "본 시스템 설계 과정에서 도출된 세 가지 요구사항은: "
        "(1) 특정 전문가만 수행 가능한 라벨링의 UI 기반 간소화, "
        "(2) 갑작스런 unknown defect 조기 탐지 [4], "
        "(3) wafer map + chip 계측값 종합 분석 UI이다. "
        "이에 데이터 생성-등록 불량 분류-미등록 불량 군집화-UI 검토를 하나로 연결하는 "
        "도메인 지식 기반 통합 파이프라인을 설계하였다.",
        indent=True, space_after=Pt(2))

    add_heading(doc, "2. 제안 방법", level=1)

    add_heading(doc, "2.1 데이터 생성 및 표현", level=2)
    add_body(doc,
        "Bucket A(Failbit 로그)와 Bucket B(FBT/QVL 계측)를 256 연결 병렬 다운로드·±10초 매칭으로 "
        "정합하고, Cython 컴파일 파싱으로 Python 대비 3~5× 속도를 달성하였다. "
        "이미지는 8-bit palette PNG(32색)로 저장한다. Failbit Map은 검사 결과로 20색 전후만 사용하므로 "
        "palette 양자화가 사실상 무손실이며, PLTE 교체만으로 사용자별 색상 scheme 즉시 전환이 가능하다. "
        "24bit RGB 대비 약 65% 크기 절감.",
        space_after=Pt(2))

    add_heading(doc, "2.2 등록 불량 분류 (16-class)", level=2)
    add_body(doc,
        "EfficientNetV2, MaxViT, ViT, Swin, ConvNeXtV2를 Hugging Face timm으로 fine-tuning 비교. "
        "총 ~1500개(val ~500개) 데이터에서 ConvNeXtV2-Base(FCMAE+IN-22k, 384×384)가 val F1 0.92로 1위. "
        "FGVC·Focal Loss는 효과 제한적. 모든 샘플에 YOLO를 적용하지 않은 이유: "
        "(1) 연산 비용, (2) wafer 불량은 전체 맵의 global 분포로 판단해야 하므로 CNN first 필수. "
        "CNN 신뢰도 < 0.80 또는 difficult class일 때만 Grad-CAM ROI [6] + spatial prior로 보강 후 "
        "YOLO 이차 검증. YOLO 학습에서 cls_loss가 가장 결정적 — ROI 내 defect class 구분이 핵심이며, "
        "cls gain 강화 시 correction precision 90% 이상, F1 0.92→0.95 달성.",
        space_after=Pt(2))

    add_heading(doc, "2.3 미등록 불량 탐지 (Open-set)", level=2)
    add_body(doc,
        "SupCon [3]은 등록 군집 밀집도는 향상시켰으나 unknown 군집 분리도 저하 → InfoNCE(SimCLR [7]) 채택. "
        "contrastive learning은 정답 예측이 아닌 임베딩 배치 학습이므로 train/test 분리 없이 "
        "전체 데이터 활용 가능. 6×6 격자(grid36_full) 기반 structured local 샘플링. "
        "flip 제외(left/right edge_loc는 공정 원인 상이). "
        "Global InfoNCE(Queue 16K) + Local InfoNCE + HDBSCAN(size≥12, prob≥0.55, persist≥0.20)으로 "
        "unknown cluster 안정화.",
        space_after=Pt(2))

    add_heading(doc, "3. 실험 결과", level=1)
    add_table(doc, TABLE1)
    add_table(doc, TABLE_BACKBONE)
    add_body(doc,
        "edge_loc·edge_ring·local 등 유사 패턴에서 ROI 보강 효과가 두드러졌으며, "
        "선택적 YOLO 보정의 correction precision 90% 이상. "
        "unknown defect 평가는 Silhouette Score, 전문가 동의율, 재학습 후 F1 향상으로 측정.",
        space_after=Pt(2))
    add_table(doc, TABLE_COMPARE)

    add_heading(doc, "4. 결론", level=1)
    add_body(doc,
        "ConvNeXtV2 + Optuna로 F1 0.92, 선택적 ROI + YOLO로 0.95, InfoNCE + HDBSCAN으로 "
        "unknown 자동 군집화, palette PNG·FBT/QVL 정합으로 multimodal 기반을 구축하였다. "
        "향후 칩 레벨 다중 라벨, 이미지+계측 multimodal 모델, 전체 제품군 확대를 예정한다.",
        indent=True, space_after=Pt(2))

    add_refs(doc, REFS[:6])
    return doc


# ══════════════════════════════════════════════════════════════
# 2페이지 Claude 버전 (paper_claude_2page_revN.docx)
# ── 초록(1단) 아래에서 바로 2단 전환 → 1페이지 내에 본문 시작
# ── 구성: 서론(FTN/QTN 원인 분석 체인) + Fig1(YOLO ROI)
#          + Fig2(UI: Failbit/FTN/QTN) + Table1(성능 3행) + 결론
# ══════════════════════════════════════════════════════════════
def build_2page_claude() -> Document:
    doc = Document()

    sec = doc.sections[0]
    sec.page_width  = Cm(21)
    sec.page_height = Cm(29.7)
    set_margins(sec)
    set_two_columns(sec, space_cm=0.5)
    doc.styles["Normal"].font.size = PT10
    doc.styles["Normal"].font.name = F_KO

    main_sp = sec._sectPr
    tp = OxmlElement("w:type")
    tp.set(qn("w:val"), "continuous")
    main_sp.insert(0, tp)

    # ── 1단: 제목 / 저자 / 초록 ──────────────────────────────
    add_title_block(doc, TITLE_KO_CLAUDE, TITLE_EN_CLAUDE)
    add_author_block(doc, AUTHORS, AFFIL)
    add_abstract_block(doc, ABSTRACT_CODEX)

    pk = doc.add_paragraph()
    rk = pk.add_run(
        "Keywords: Failbit Map, Wafer Defect Analysis, ConvNeXtV2, "
        "Grad-CAM, YOLO, Contrastive Learning, HDBSCAN, Multimodal"
    )
    _apply_run_font(rk, size=PT9, italic=True)
    _set_single_spacing(pk)
    pk.paragraph_format.space_after = Pt(4)

    add_1col_header_break(doc)

    # ── 2단: 본문 ─────────────────────────────────────────────

    # 1. 서론 (codex 5-paragraph structure)
    add_heading(doc, "1. 서론", level=1)
    add_body(doc,
        "Failbit Map은 EDS(Electrical Die Sorting) Test에서 각 Memory Cell Block의 불량 cell 개수를 "
        "Grade 0(정상)부터 7(최대 불량)까지로 표현한 데이터이다. Wafer 1장에는 약 1,000만 개의 block이 "
        "포함되므로, Failbit Map은 불량의 위치, 밀집도, 형상, 방향성을 반영하는 초고밀도 공간 데이터이자 "
        "수율 저하 원인 분석의 핵심 데이터이다. 그러나 현업 분석은 주로 Chip 단위 Measure(측정 항목별 집계값)의 "
        "발생 빈도를 기준으로 이상 여부를 판단하고 있어, Failbit Map 전체에서만 드러나는 불량 분포의 공간 패턴을 "
        "충분히 포착하지 못하며 수율 저하 원인 추적에 사각지대가 발생할 수 있다. 따라서 정밀한 원인 분석을 "
        "위해서는 Measure 기반 접근만으로는 충분하지 않으며, Failbit Map 자체를 핵심 분석 단위로 삼아야 한다.",
        indent=True, space_after=Pt(2))
    add_body(doc,
        "실제 현업 적용에는 두 가지 제약이 있다. 첫째, 대량 Map 생성이 어렵다. 설비 Log는 Wafer당 10~50MB "
        "수준이며 특정 제품군에서는 하루 약 2,000장의 Wafer가 발생하지만, 기존 분석 환경에서는 속도와 메모리 "
        "제약으로 한 번에 48매 이하만 처리할 수 있다. 둘째, Map이 생성되더라도 불량 여부와 유형 판단은 "
        "여전히 엔지니어의 직접 확인에 의존하고 있어 전수 분석이 어렵다. 본 논문은 대량 Raw Data를 처리하는 "
        "데이터 파이프라인과 Failbit Map 기반 Known 불량 분류 및 Unknown 불량 검출을 통합한 AI 아키텍처를 "
        "구현하여 실데이터에 적용·검증하였다. 본 논문의 주요 기여는 다음과 같다.",
        indent=True, space_after=Pt(2))
    add_body(doc,
        "첫째, Test 부서와 협의하여 대량 Raw Data 적재 체계와 전수 Failbit Map 생성 파이프라인을 구축하였다. "
        "palette-indexed PNG를 적용해 이미지 용량을 약 75% 절감하고, 실시간 설비 Log 수집과 1시간 주기 "
        "Failbit Map 생성 체계를 구현하여 대량 분석이 가능한 운영 기반을 마련하였다.",
        indent=True, space_after=Pt(2))
    add_body(doc,
        "둘째, Known 불량의 정밀 분류를 위해 ConvNeXt V2와 ROI 기반 YOLO를 결합한 2단계 분석 구조를 "
        "제안하였다. ConvNeXt V2로 Wafer 전역 패턴을 우선 분류하고, 신뢰도가 낮은 샘플에 한해 ROI 기반 "
        "YOLO로 불량 영역의 위치 정보를 추가 반영함으로써 전역 패턴과 국소 위치 정보를 함께 활용하는 정밀 "
        "분류를 구현하였다.",
        indent=True, space_after=Pt(2))
    add_body(doc,
        "셋째, 기존 분류 체계에 등록되지 않은 상태에서 새롭게 발생하는 Unknown 불량을 검출하기 위해 "
        "SimCLR 계열 contrastive learning 기반 분석 구조를 구현하였다. grid structured local sampling을 "
        "적용하여 동일 형태의 불량이라도 발생 위치에 따라 구분될 수 있도록 함으로써, Unknown 불량을 "
        "유사 패턴 그룹 수준에서 검출·해석할 수 있도록 하였다.",
        indent=True, space_after=Pt(2))

    # 2. 제안 방법
    add_heading(doc, "2. 제안 방법", level=1)

    add_subheading(doc, "2.1 데이터 파이프라인")
    add_body(doc,
        "설비 Log를 Storage에 실시간 적재하고 1시간 주기로 Failbit Map으로 변환·생성한다. 핵심 병목인 "
        "raw hex→grade 변환(웨이퍼당 약 1,000만 pixel 반복)을 Cython으로 전환하여 약 100배 가속하였다. "
        "i번째 입력 Hex 문자열을 h_i라 할 때, 십진수 변환값 v_i와 최종 Grade g_i의 디코딩 수식은 다음과 같다.",
        space_after=Pt(1))
    add_equation(doc, "v_i = int(h_i, 16)", number="1")
    add_equation(doc, "g_i = max(v_i − 9, 0)", number="2")
    add_body(doc,
        "Failbit Map은 grade·boundary 등 제한된 색상만 사용하므로 RGB PNG 대신 palette-indexed PNG를 "
        "적용하여 RGB 대비 약 75% 용량을 절감하였다. 32-색 팔레트 P를 정의하고 각 픽셀을 팔레트 인덱스로 "
        "치환하며, 좌표 (x,y)에서의 RGB 값은 인덱스 맵 I_idx를 통해 다음과 같이 매핑된다.",
        space_after=Pt(1))
    add_equation(doc,
                 "I_rgb(x,y) = P[I_idx(x,y)],   I_idx(x,y) ∈ {0,1,…,31}",
                 number="3")
    add_body(doc,
        "Measure 로그와 wafer ID 기준 매칭으로 chip별 FTN·QTN·BIN을 annotation에 사전 적재한다.",
        space_after=Pt(2))

    add_subheading(doc, "2.2 Known 불량 분류 (16-class)")
    add_body(doc,
        "약 1,500장(train/test 4:1 층화 분할)의 라벨 데이터로 ConvNeXtV2-Base [5](FCMAE+IN-22k, 384×384)를 "
        "backbone으로 채택하고 Optuna HPO [3]로 최적화하였다. center·edge·local 계열처럼 wafer 전역 형상이 "
        "유사한 클래스 쌍은 전역 문맥만으로 분리가 어려워 confidence가 낮게 출력되므로, "
        "confidence < 0.80 샘플에 한해 Grad-CAM ROI [6]를 추출하고 YOLO [9]가 chip 단위 pixel 패턴을 "
        "직접 검증하여 오분류를 보정한다(Fig. 1).",
        space_after=Pt(2))

    _fig1_path = str(OUT_DIR / "_fig_yolo_roi_v7.png")
    make_yolo_roi_figure(_fig1_path)
    add_figure(doc, _fig1_path,
               "Fig. 1. wafer 전역(chip 배치)은 A·B 동일하여 CNN이 혼동되나, chip 내 pixel 패턴(star/x)으로 YOLO가 교정.",
               width_cm=5.8)

    add_subheading(doc, "2.3 Unknown 불량 탐지 (Open-set)")
    add_body(doc,
        "Unknown 불량은 사전에 정답 클래스를 정의하기 어려운 실제 운영 이미지를 대상으로 하므로, "
        "이를 후보 grouping 문제로 정의하였다. 5일치 운영 데이터 약 10,000장으로 SimCLR [7] 계열 "
        "contrastive learning 기반 임베딩을 학습하고, 1일치 2,000장에 HDBSCAN [8](size≥12, prob≥0.55)을 "
        "적용하여 유사 패턴을 그룹 단위로 압축한다. 이때 wafer의 zone 기반 불량 해석 특성을 반영하기 위해 "
        "6×6 grid structured local sampling을 적용한다.",
        space_after=Pt(2))

    # 3. 실험 결과
    add_heading(doc, "3. 실험 결과", level=1)
    add_table(doc, TABLE_PERF_CLAUDE)
    add_body(doc,
        "하루 약 2,000장의 Wafer가 1시간 주기로 발생하는 운영 제약 하에서 backbone 선택은 F1과 추론 처리량을 "
        "함께 고려하였다. 각 계열의 Base/M급(약 70~100M params)만 비교하고 Large 이상은 제외하였으며, "
        "MaxViT-Base와 ConvNeXtV2-Base는 F1 0.87로 동등하나 추론 속도 우위로 ConvNeXtV2-Base를 채택하였다. "
        "Baseline 0.78 대비 Optuna HPO로 F1 0.92, 선택적 ROI-YOLO 보정으로 최종 0.95를 달성하였다. "
        "Unknown 탐지에서는 하루 2,000장이 13개 후보 그룹으로 자동 압축되어(약 154배 축소) "
        "엔지니어 검토 결과 7개(53.8%)는 신규 불량, 6개는 1 lot 헌팅 또는 이미지 이상으로 정리되어 "
        "후보 리스트 전체가 검토 가능한 작업 단위로 기능하였다.",
        space_after=Pt(2))

    # 4. 결론
    add_heading(doc, "4. 결론", level=1)
    add_body(doc,
        "본 연구는 데이터 파이프라인과 Known·Unknown 통합 AI 아키텍처로 Failbit Map을 수작업 시각 자료에서 "
        "대량 운영 데이터 기반 핵심 분석 단위로 전환하였다. 현재 DRAM 생산 라인에서 운영 중이며, 향후 "
        "전 제품군 확산과 FTN·QTN 기반 multimodal 융합으로 확장할 계획이다.",
        indent=True, space_after=Pt(2))

    add_refs(doc, REFS_CLAUDE)
    return doc


def build_codex_revised() -> Document:
    doc = Document()

    sec = doc.sections[0]
    sec.page_width = Cm(21)
    sec.page_height = Cm(29.7)
    set_margins(sec)
    set_two_columns(sec, space_cm=0.5)
    doc.styles["Normal"].font.size = PT10
    doc.styles["Normal"].font.name = F_KO

    main_sp = sec._sectPr
    tp = OxmlElement("w:type")
    tp.set(qn("w:val"), "continuous")
    main_sp.insert(0, tp)

    add_title_block(doc, TITLE_KO_CODEX, TITLE_EN_CODEX)
    add_author_block(doc, AUTHORS, AFFIL)
    add_abstract_block(doc, ABSTRACT_CODEX)

    pk = doc.add_paragraph()
    rk = pk.add_run(
        "Keywords: Failbit Map, Wafer Defect Analysis, ConvNeXtV2, "
        "YOLO, Contrastive Learning, HDBSCAN"
    )
    _apply_run_font(rk, size=PT9, italic=True)
    _set_single_spacing(pk)
    pk.paragraph_format.space_after = Pt(4)

    add_1col_header_break(doc)

    add_heading(doc, "1. 서론", level=1)
    add_body(doc,
        "Failbit Map은 EDS Test에서 Memory Cell Block 단위의 불량 정도를 Grade 0(정상)부터 7(최대 불량)까지로 표현한 데이터이다. "
        "Wafer 1장에는 약 1,000만 개의 block이 존재하므로, Failbit Map은 불량의 위치와 형태를 반영하는 초고해상도 데이터이자 수율 저하 원인 분석의 핵심 분석 대상이다. "
        "그러나 현업 분석은 Wafer 내 약 1,000개 Chip에서 산출된 Measure 값의 발생 개수를 기반으로 이상 여부를 판단하고 있어, "
        "Failbit Map에서만 발현되는 불량을 검출하지 못한다. 따라서 수율 개선 및 Drop 방지를 위해서는 Measure 기반 접근만으로는 충분하지 않으며, "
        "Failbit Map 자체를 핵심 분석 단위로 삼아야 한다 [1]-[3].",
        indent=True, space_after=Pt(2))
    add_body(doc,
        "실제 현업 적용에는 두 가지 제약이 있다. 첫째, 대량 Map 생성이 어렵다. 설비 Log는 Wafer당 10~50MB 수준이며, 특정 제품에서는 하루 약 2,000장의 Wafer가 발생한다. "
        "그러나 기존 환경은 속도와 메모리 제약으로 대량 처리가 어려웠으며, 실제로는 한 번에 48매까지 확인 가능하여 전수 분석에 한계가 있었다. "
        "둘째, Map이 생성되더라도 불량 여부와 유형 판단이 엔지니어의 수동 판독에 의존하므로 전수 분석이 어렵다. "
        "본 논문은 이러한 한계를 해결하기 위해, 대량 Raw Data 처리를 위한 데이터 파이프라인과 Failbit Map 기반 Known 불량 분류, Unknown 불량 검출을 통합한 분석 아키텍처를 제안한다. "
        "주요 기여는 다음과 같다.",
        indent=True, space_after=Pt(2))
    add_body(doc,
        "첫째, 대량 설비 Log의 실시간 적재와 1시간 주기 Failbit Map 생성 체계를 구현하여, 대량 Map의 지속적 생성과 운영 활용이 가능한 데이터 처리 기반을 구축하였다. "
        "둘째, Known 불량 분류는 ConvNeXt V2로 Wafer 내 불량 Chip들의 종류와 분포 패턴을 1차 분류하고, 유사 분포로 인해 혼동되는 샘플은 ROI 기반 YOLO로 개별 Chip object detection을 수행하여 "
        "2차 분류함으로써 성능을 향상시켰다. 셋째, 기존 분류 체계에 등록되지 않은 상태에서 새롭게 발생하는 Unknown 불량을 검출하기 위해 SimCLR 계열 contrastive learning 기반 분석 구조를 구현하였다. "
        "또한 grid structured local sampling을 적용하여 발생 위치까지 반영함으로써, Unknown 불량을 보다 정밀하게 검출할 수 있도록 하였다.",
        indent=True, space_after=Pt(2))

    add_heading(doc, "2. 제안 방법 및 결과", level=1)

    add_subheading(doc, "2.1 데이터 파이프라인")
    add_body(doc,
        "주요 병목은 wafer당 약 1,000만 개의 암호화된 test 결과를 grade 값으로 변환하는 처리 속도와, 4K를 초과하는 초고해상도 이미지의 저장 용량 부담이었다. "
        "이에 Cython 최적화로 데이터 변환 속도를 약 100배 향상시켰으며(Fig. 1), palette-indexed PNG 적용으로 이미지 용량을 약 75% 절감하였다(Fig. 2).",
        space_after=Pt(2))
    add_labeled_example_table(
        doc,
        "Hex-to-grade conversion",
        [
            ("Raw:", "090B0C0D0E0F090A0B0C"),
            ("Decoding:", "\"0C\" -> \"C\" -> 12 (hex to decimal) -> 3 (if value != 0, subtract 9)"),
            ("Python:", "interpreter-based loop execution"),
            ("Cython:", "compiled integer loop execution"),
            ("Grade:", "0 2 3 4 5 6 0 1 2 3"),
        ],
        "Fig. 1. Hex-to-grade conversion accelerated by Cython.",
    )
    add_labeled_example_table(
        doc,
        "RGB PNG vs Palette-indexed PNG",
        [
            ("RGB PNG:", [
                "[(123,54,24), (123,54,24), ..., (123,54,24)],",
                "[(123,54,24), (123,54,24), ..., (123,54,24)]",
            ]),
            ("Palette-indexed PNG:", [
                "P[3] = (123,54,24)",
                "[(3), (3), ..., (3)],",
                "[(3), (3), ..., (3)]",
            ]),
            ("Result:", "(123,54,24) -> (3)"),
        ],
        "Fig. 2. Palette-indexed PNG for failbit map compression.",
    )
    add_body(doc,
        "이러한 파이프라인으로 생성된 경량화 Failbit Map은 다음의 Known 불량 및 Unknown 불량 분석 AI 입력으로 사용된다.",
        space_after=Pt(2))

    add_subheading(doc, "2.2 Known 불량 분류")
    add_body(doc,
        "Known 불량 분석은 16개 기등록 클래스를 대상으로 하며, 1,500개의 라벨된 Failbit Map으로 "
        "ConvNeXtV2[4] 기반 1단계 wafer-level 분류기와 저신뢰 샘플 대상 2단계 ROI 기반 YOLO를 구성하였다.",
        space_after=Pt(2))
    add_body(doc,
        "혼동되는 클래스는 wafer 내 defect chip distribution은 유사하지만 defect chip 내부 pixel morphology가 다르므로, "
        "2단계 ROI-YOLO가 국소 pattern을 추가로 판별하도록 구성하였다(Fig. 3).",
        space_after=Pt(2))

    _fig_known_path = str(OUT_DIR / "_fig_yolo_roi.png")
    make_yolo_roi_figure(_fig_known_path)
    add_figure(
        doc,
        _fig_known_path,
        "Fig. 3. Two-stage known-defect classification with ROI-YOLO.",
        width_cm=8.2,
    )
    add_table(doc, TABLE_PERF_CLAUDE)
    add_body(doc,
        "전체 운영 기준 하루 약 2만 장 이상의 Wafer가 1시간 주기로 처리되므로, backbone 선택에서는 정확도와 추론 처리량을 함께 고려하였다. "
        "MaxViT[7]와 ConvNeXtV2 (Ref)는 동일한 test Weighted F1 0.87을 보였으나, ConvNeXtV2가 1시간 주기 처리 요건을 더 안정적으로 만족할 수 있어 최종 backbone으로 선정하였다. "
        "Baseline CNN의 test Weighted F1은 0.78에서 0.87(Ref), 0.92(Ref + Optuna), 0.95(Ref + Optuna + ROI)로 단계적으로 향상되었다.",
        space_after=Pt(2))

    add_subheading(doc, "2.3 Unknown 불량 검출")
    add_body(doc,
        "Unknown 불량은 후보 grouping 문제로 정의하였다. 5일치 운영 데이터 10,000장으로 SimCLR 계열 "
        "contrastive learning[5] 기반 임베딩을 학습하고, 별도 1일치 2,000장에 HDBSCAN[6]을 적용하여 유사 패턴을 그룹화하였다. "
        "또한 Wafer 이미지를 N×N grid로 균등 분할하고 동일 grid cell 내 샘플을 positive pair로 구성하는 "
        "grid structured local sampling으로 발생 위치 정보를 반영하였다.",
        space_after=Pt(2))

    _fig_unknown_path = str(OUT_DIR / "_fig_cluster.png")
    add_figure(
        doc,
        _fig_unknown_path,
        "Fig. 4. Unknown-defect grouping on production images.",
        width_cm=8.2,
    )
    add_body(doc,
        "Unknown 경로에서는 운영 이미지 2,000장에 대한 grouping 결과 13개 후보 그룹이 검출되었으며, 현업 분석 엔지니어 검증 결과 이 중 7개가 실제 불량 그룹으로 판정되었다. "
        "나머지 6개 그룹은 lot성 warning 수준의 noise이거나 chip 불량으로 이어지지 않는 패턴으로 해석되었다.",
        space_after=Pt(2))

    add_heading(doc, "3. 결론", level=1)
    add_body(doc,
        "본 연구는 수율 개선 업무의 핵심인 Failbit Map 분석을 위해 대량 생성 파이프라인과 "
        "Known 불량 및 Unknown 불량 분석 AI를 통합함으로써, 기존 수작업 중심 분석을 현업 적용 가능한 AI 자동화 체계로 고도화하였다. "
        "현재 DRAM 생산 라인에서 운영 중이며, 1시간 주기 Map 생성과 등록·미등록 불량 자동 분석을 통해 FAB 품질 분석과 신규 불량 대응에 실질적으로 기여하고 있다.",
        indent=True, space_after=Pt(2))

    add_refs(doc, REFS_CODEX)
    return doc


# ══════════════════════════════════════════════════════════════
# 실행
# ══════════════════════════════════════════════════════════════
if __name__ == "__main__":
    print("Building 2-page claude version...")
    doc_cl = build_2page_claude()
    existing_cl = []
    for search_dir in [OUT_DIR, OUT_DIR / "archive"]:
        for p in search_dir.glob("paper_claude_2page_rev*.docx"):
            try:
                existing_cl.append(int(p.stem.rsplit("rev", 1)[1]))
            except (IndexError, ValueError):
                pass
    rev_cl = (max(existing_cl) + 1) if existing_cl else 1
    path_cl = OUT_DIR / f"paper_claude_2page_rev{rev_cl}.docx"
    doc_cl.save(str(path_cl))
    print(f"  -> saved: {path_cl}")
